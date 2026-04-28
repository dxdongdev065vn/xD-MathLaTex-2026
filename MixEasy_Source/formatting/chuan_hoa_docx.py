import os

import tkinter as tk

from tkinter import ttk, Label, filedialog, Entry, Button, Frame, Listbox, Scrollbar, messagebox

from tkinter import simpledialog

from spire.doc import Section, Paragraph, Table, TextRange, Shape, FileFormat, HorizontalAlignment, Regex

from spire.doc import UnderlineStyle

from spire.doc import Document as SpireDocument

from spire.doc import Color

from spire.doc import DocumentObjectType

from spire.doc import OfficeMath

from docx import Document

from docx.text.paragraph import Paragraph

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import RGBColor

from docx.shared import Inches, Pt

from docx.oxml import parse_xml

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from lxml import etree



ElementTree

from xml.etree.ElementTree import QName

import xml.etree.ElementTree, etree

import re

import zipfile

import bisect

from tools import tool_by_docx as docxtool

from mixing import mix_docx_func as mdocx_f



def InchesToPoints(inches):

    return inches * 72





def LinesToPoints(lines):

    return lines * 12





def open_doc_off_python_docx(doc_path):

    doc = Document(doc_path)

    return doc





def remove_last_element(lst):

    return lst[:-1]





def resize_images_to_column_width(doc):

    '''Chỉnh kích thước hình ảnh sao cho không quá 3.5 inches chiều rộng'''

    max_width = Inches(3.5)

    for shape in doc.inline_shapes:

        if not shape.width > max_width:

            continue

        aspect_ratio = shape.height / shape.width

        shape.width = int(max_width)

        shape.height = int(max_width * aspect_ratio)





def page_2_cot_Mix(doc):

    '''Thiết lập trang thành 2 cột với khoảng cách phù hợp'''

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.4)

    section.bottom_margin = Inches(0.4)

    section.left_margin = Inches(0.4)

    section.right_margin = Inches(0.4)

    section.gutter = Inches(0)

    section.header_distance = Inches(0.24)

    section.footer_distance = Inches(0.24)

    sectPr = doc.sections[0]._sectPr

    cols = OxmlElement('w:cols')

    cols.set(qn('w:num'), '2')

    cols.set(qn('w:sep'), '1')

    spacing_in_twips = int(144)

    cols.set(qn('w:space'), str(spacing_in_twips))

    sectPr.append(cols)





def delete_bookmarks(doc):

    excluded_names = {

        'MDH',

        'num_page'}

    bookmark_start_ids_to_keep = set()

    for elem in doc.element.body.iter():

        if not elem.tag == qn('w:bookmarkStart'):

            continue

        name = elem.get(qn('w:name'))

        bookmark_id = elem.get(qn('w:id'))

        if not name in excluded_names:

            continue

        bookmark_start_ids_to_keep.add(bookmark_id)

# WARNING: Decompyle incomplete





def delete_bookmarks_OLD(doc):

    '''Xóa tất cả bookmark trừ những bookmark được giữ lại'''

    excluded_bookmarks = {

        'MDH',

        'MDH2',

        'num_page'}

    body = doc.element.body

    bookmarks_to_remove = []

    for elem in body.iter():

        if not elem.tag.endswith('bookmarkStart'):

            continue

        name = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')

        if not name:

            continue

        if not name not in excluded_bookmarks:

            continue

        bookmarks_to_remove.append(elem)

    for bookmark in bookmarks_to_remove:

        bookmark.getparent().remove(bookmark)





def Danh_STT_Cau(doc, indices, i):

    '''Thay số thứ tự câu hỏi bằng cách xóa `run` cũ và thêm `Câu i.` vào đầu'''

    pattern = '^(Câu )([0-9]{1,})'

    body = doc.element.body

    elements = list(body)

    for p in indices:

        para = elements[p]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        runs = paragraph.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            full_text += run.text

            match_run_indices.append(idx)

            if not re.match(pattern, full_text.strip()):

                continue

            enumerate(runs)

    match = re.match(pattern, full_text.strip())

    if not re.match(pattern, full_text.strip()):

        continue

    remaining = full_text[match.end():]

    for idx in reversed(match_run_indices[1:]):

        paragraph._element.remove(runs[idx]._element)

    runs[0].text = f'''Câu {i}{remaining}'''

    runs[0].bold = True

    runs[0].font.color.rgb = RGBColor(0, 0, 255)

    i += 1

    continue





def Danh_STT_Cau_TL(doc, indices, i):

    '''Thay số thứ tự câu hỏi bằng cách xóa `run` cũ và thêm `Câu i.` vào đầu'''

    pattern = '^(Câu )([0-9]{1,})'

    body = doc.element.body

    elements = list(body)

    for p in indices:

        para = elements[p]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        runs = paragraph.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            full_text += run.text

            match_run_indices.append(idx)

            if not re.match(pattern, full_text.strip()):

                continue

            enumerate(runs)

    match = re.match(pattern, full_text.strip())

    if not match:

        continue

    remaining = full_text[match.end():]

    for idx in reversed(match_run_indices[1:]):

        paragraph._element.remove(runs[idx]._element)

    runs[0].text = f'''Câu {i}{remaining}'''

    runs[0].bold = True

    runs[0].font.color.rgb = RGBColor(0, 0, 255)

    i += 1

    continue





def Danh_STT_Cau_TL_All(doc):

    all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, 'S4@', 'E4@')

    if len(all_indices_cau) > 0:

        i = 1

        for idx in range(0, len(all_indices_cau)):

            cau = all_indices_cau[idx]

            Danh_STT_Cau_TL(doc, cau, i)

            i += len(cau) - 1

        return None





def Danh_STT_Cau_All(doc):

    for k in range(1, 5):

        all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, f'''S{k}@''', f'''E{k}@''')

        if not len(all_indices_cau) > 0:

            continue

        i = 1

        for idx in range(0, len(all_indices_cau)):

            cau = all_indices_cau[idx]

            Danh_STT_Cau(doc, cau, i)

            i += len(cau) - 1





def Danh_STT_Cau_EN(doc, indices, i):

    '''Thay số thứ tự câu hỏi tiếng Anh & tiếng Việt, giữ nguyên định dạng'''

    pattern1 = '^Question [0-9]{1,}'

    pattern2 = '^Câu [0-9]{1,}'

    body = doc.element.body

    elements = list(body)

    for p in indices:

        para = elements[p]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        runs = paragraph.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            full_text += run.text

            match_run_indices.append(idx)

            if not re.match(pattern1, full_text.strip()) and re.match(pattern2, full_text.strip()):

                continue

            enumerate(runs)

    match = re.match(pattern1, full_text.strip())

    if re.match(pattern1, full_text.strip()):

        remaining = full_text[match.end():]

        for idx in reversed(match_run_indices[1:]):

            paragraph._element.remove(runs[idx]._element)

        runs[0].text = f'''Question {i}{remaining}'''

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(0, 0, 255)

        i += 1

    match = re.match(pattern2, full_text.strip())

    if not re.match(pattern2, full_text.strip()):

        continue

    remaining = full_text[match.end():]

    for idx in reversed(match_run_indices[1:]):

        paragraph._element.remove(runs[idx]._element)

    runs[0].text = f'''Câu {i}{remaining}'''

    runs[0].bold = True

    runs[0].font.color.rgb = RGBColor(0, 0, 255)

    i += 1

    continue





def Danh_STT_Cau_All_EN(doc):

    all_indices_cau = mdocx_f.tim_cau_trong_phan_EN(doc)

    if len(all_indices_cau) > 0:

        i = 1

        for idx in range(0, len(all_indices_cau)):

            cau = all_indices_cau[idx]

            Danh_STT_Cau_EN(doc, cau, i)

            i += len(cau) - 1

        return None





def STT_ABCD_trong_cac_vung(doc, text_start, text_end):

    '''Cập nhật nhãn A, B, C, D trong các vùng được chỉ định (chỉ thay ký tự đầu của run[0])'''

    labels = [

        'A',

        'B',

        'C',

        'D']

    all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, text_start, text_end)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            phuong_an = phuong_an[:-1]

            encounter_count = 0

            for p in phuong_an:

                para = paragraphs[p]

                if not para.tag.endswith('p'):

                    continue

                paragraph = Paragraph(para, doc)

                runs = paragraph.runs

                if not runs:

                    continue

                first_run = runs[0]

                text = first_run.text.strip()

                if not text:

                    continue

                if not text[0] in 'ABCD':

                    continue

                current_label = labels[encounter_count % len(labels)]

                first_run.text = current_label + first_run.text[1:]

                encounter_count += 1





def STT_ABCD_trong_cac_vung_EN(doc):

    '''Cập nhật nhãn A, B, C, D trong các vùng được chỉ định (chỉ thay ký tự đầu của run[0])'''

    labels = [

        'A',

        'B',

        'C',

        'D']

    all_indices_cau = mdocx_f.tim_cau_trong_phan_EN(doc)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I_EN(doc, cau[k], cau[k + 1])

            phuong_an = phuong_an[:-1]

            encounter_count = 0

            for p in phuong_an:

                para = paragraphs[p]

                if not para.tag.endswith('p'):

                    continue

                paragraph = Paragraph(para, doc)

                runs = paragraph.runs

                if not runs:

                    continue

                first_run = runs[0]

                text = first_run.text.strip()

                if not text:

                    continue

                if not text[0] in 'ABCD':

                    continue

                current_label = labels[encounter_count % len(labels)]

                first_run.text = current_label + first_run.text[1:]

                encounter_count += 1





def STT_ABCD(doc):

    STT_ABCD_trong_cac_vung(doc, 'S1@', 'E1@')





def STT_ABCD_EN(doc):

    STT_ABCD_trong_cac_vung_EN(doc)





def STT_abcd_nho(doc):

    labels = [

        'a',

        'b',

        'c',

        'd']

    all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, 'S2@', 'E2@')

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_II(doc, cau[k], cau[k + 1])

            phuong_an = phuong_an[:-1]

            encounter_count = 0

            for p in phuong_an:

                para = paragraphs[p]

                if not para.tag.endswith('p'):

                    continue

                paragraph = Paragraph(para, doc)

                runs = paragraph.runs

                if not runs:

                    continue

                first_run = runs[0]

                text = first_run.text.strip()

                if not text:

                    continue

                if not text[0] in 'abcd':

                    continue

                current_label = labels[encounter_count % len(labels)]

                first_run.text = current_label + first_run.text[1:]

                encounter_count += 1





def tim_thang_trong_all(doc):

    '''Tìm tất cả vị trí chứa `#X` trong tài liệu'''

    indices_cau = []

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, para in enumerate(paragraphs):

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        if not re.search('#(\\d+)', paragraph.text):

            continue

        indices_cau.append(i)

    return indices_cau





def tim_cau_trong_all(doc):

    """Tìm tất cả vị trí chứa 'Question X.' hoặc 'Câu X.' trong tài liệu"""

    pattern1 = '^Question [0-9]{1,}[.:]'

    pattern2 = '^Câu [0-9]{1,}[.:]'

    indices_cau = []

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, para in enumerate(paragraphs):

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        if not re.match(pattern1, paragraph.text) and re.match(pattern2, paragraph.text):

            continue

        indices_cau.append(i)

    return indices_cau





def STT_thang_OK(doc):

    '''Cập nhật số thứ tự `#X` dựa trên giá trị X, đảm bảo đúng thứ tự'''

    a = tim_cau_trong_all(doc)

    b = tim_thang_trong_all(doc)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for bb in b:

        para = paragraphs[bb]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        matches = re.findall('#(\\d+)', paragraph.text)

        if not matches:

            continue

        full_text = (lambda .0: pass# WARNING: Decompyle incomplete

)(paragraph.runs())

        for match in matches:

            first_question_index = bisect.bisect_left(a, bb)

            X_value = int(match)

            new_index = first_question_index + X_value

            full_text = re.sub(f'''#{match}''', str(new_index), full_text)

        offset = 0

        for run in paragraph.runs:

            length = len(run.text)

            run.text = full_text[offset:offset + length]

            offset += length





def STT_thang(doc):

    '''Cập nhật số thứ tự `#X`, xử lý cả trường hợp `#X` bị tách bởi MathType'''

    a = tim_cau_trong_all(doc)

    b = tim_thang_trong_all(doc)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for bb in b:

        para = paragraphs[bb]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        elements = list(para.iter())

        full_text = ''

        text_elements = []

        for element in elements:

            if not element.tag.endswith('t'):

                continue

            if not element.text:

                continue

            text_elements.append(element)

            full_text += element.text

        matches = re.findall('#(\\d+)', full_text)

        if not matches:

            continue

        for match in matches:

            first_question_index = bisect.bisect_left(a, bb)

            X_value = int(match)

            new_index = first_question_index + X_value

            full_text = re.sub(f'''#{match}''', str(new_index), full_text)

        offset = 0

        for element in text_elements:

            length = len(element.text)

            element.text = full_text[offset:offset + length]

            offset += length





def canh_before_after(doc, p):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_4_at(doc, p):

    tab_stops = [

        Inches(0.21),

        Inches(1.97),

        Inches(3.73),

        Inches(5.49)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_at(doc, p):

    tab_stops = [

        Inches(0.21),

        Inches(3.73)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_cot_4_at(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(0.98),

        Inches(1.86),

        Inches(2.74)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_cot_2_at(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(1.86)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def insert_tab_before_first_child(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(1.86)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def insert_tab_before_indices(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)





def MergeChildObjectsIntoGroup(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_1(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_1_2_cot(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_2_1_1(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_2_1_1_2_cot(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_4(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)

        tab_btp_4_at(doc, p)





def option_4_2_cot(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)

        tab_btp_2_cot_4_at(doc, p)





def len_option(doc, p):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    if p < 0 or p >= len(paragraphs):

        raise IndexError('Chỉ số p không hợp lệ.')

    para = paragraphs[p]

    if not para.tag.endswith('p'):

        return 0

    text_length = 0

    equation_length = 0

    mathtype_length = 0

    image_length = 0

    nsmap = doc.element.nsmap.copy()

    if 'a' not in nsmap:

        nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    if 'm' not in nsmap:

        nsmap['m'] = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    if 'o' not in nsmap:

        nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

    for run in para.findall('.//w:r', namespaces = nsmap):

        for t in run.findall('.//w:t', namespaces = nsmap):

            if not t.text:

                continue

            text_length += len(t.text)

    math_objects = para.findall('.//m:oMathPara', namespaces = nsmap)

    if not math_objects:

        math_objects = para.findall('.//m:oMath', namespaces = nsmap)

    for math in math_objects:

        for mt in math.findall('.//m:t', namespaces = nsmap):

            if not mt.text:

                continue

            equation_length += len(mt.text)

    equation_length = int(1.7 * equation_length)

# WARNING: Decompyle incomplete





def max_len_option(doc, indices):

    max_length = 0

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for p in indices:

        if p < 0 or p >= len(paragraphs):

            raise IndexError('Chỉ số p không hợp lệ.')

        current_length = len_option(doc, p)

        if not current_length > max_length:

            continue

        max_length = current_length

    return max_length





def tim_cau_trong_phan(doc, text_start, text_end):

    (indices_S, indices_E) = mdocx_f.find_phan(doc, text_start, text_end)

    all_indices_cau = []

    if len(indices_S) != len(indices_E):

        messagebox.showinfo('Thông báo', f'''Số kí hiệu {text_start} và {text_end} không bằng nhau, hãy chỉnh sửa phù hợp''')

        return None

    if len(indices_S) > 0:

        pattern1 = '^Câu [0-9]{1,}[.:]'

        pattern2 = '^Question [0-9]{1,}[.:]'

        pattern3 = '@'

        body = doc.element.body

        paragraphs = list(body.iterchildren())

        for start, end in zip(indices_S, indices_E):

            indices_cau = []

            for j in range(start, end):

                para = paragraphs[j]

                text = (lambda .0: pass# WARNING: Decompyle incomplete

)(para.iter()()).strip()

                if not re.match(pattern1, text) and re.match(pattern2, text) and re.search(pattern3, text):

                    continue

                indices_cau.append(j)

            indices_cau.append(end)

            all_indices_cau.append(indices_cau)

    return all_indices_cau





def tim_cau_first(doc):

    """Tìm chỉ mục đoạn văn đầu tiên chứa 'Câu X.' hoặc 'Câu X:'"""

    pattern1 = '^Câu [0-9]{1,}[.:]'

    pattern2 = '^Question [0-9]{1,}[.:]'

    for index, paragraph in enumerate(doc.paragraphs):

        if not re.match(pattern1, paragraph.text) and re.match(pattern2, paragraph.text):

            continue

        

        return enumerate(doc.paragraphs), index

    return 0





def lay_size_font(doc, index):

    '''Lấy kích thước font của văn bản đầu tiên trong đoạn văn tại index.'''

    if index < len(doc.paragraphs):

        paragraph = doc.paragraphs[index]

        for run in paragraph.runs:

            if not run.font:

                continue

            if not run.font.size:

                continue

            

            return paragraph.runs, run.font.size.pt

    return 12





def sap_xep_phan_I_trong_cac_vung(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 25:

                option_1(doc, phuong_an)

                continue

            if L1max < 49:

                option_2_1_1(doc, phuong_an)

                continue

            option_4(doc, phuong_an)





def sap_xep_phan_I_trong_cac_vung_font13(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 23:

                option_1(doc, phuong_an)

                continue

            if L1max < 45:

                option_2_1_1(doc, phuong_an)

                continue

            option_4(doc, phuong_an)





def sap_xep_phan_I_trong_cac_vung_font14(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 21:

                option_1(doc, phuong_an)

                continue

            if L1max < 41:

                option_2_1_1(doc, phuong_an)

                continue

            option_4(doc, phuong_an)





def sap_xep_phan_I_trong_cac_vung_Japan(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 9:

                option_1(doc, phuong_an)

                continue

            if L1max < 18:

                option_2_1_1(doc, phuong_an)

                continue

            option_4(doc, phuong_an)





def sap_xep_phan_I(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_trong_cac_vung_font13(doc, 'S1@', 'E1@')

        return None

    if size_font == 14:

        sap_xep_phan_I_trong_cac_vung_font14(doc, 'S1@', 'E1@')

        return None

    sap_xep_phan_I_trong_cac_vung(doc, 'S1@', 'E1@')





def sap_xep_phan_I_EN(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_trong_cac_vung_font13(doc, '<S@>', '<E@>')

        return None

    if size_font == 14:

        sap_xep_phan_I_trong_cac_vung_font14(doc, '<S@>', '<E@>')

        return None

    sap_xep_phan_I_trong_cac_vung(doc, '<S@>', '<E@>')





def sap_xep_phan_I_Japan(doc):

    sap_xep_phan_I_trong_cac_vung_Japan(doc, '<S@>', '<E@>')





def sap_xep_phan_I_2_cot_trong_cac_vung(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 12:

                option_1_2_cot(doc, phuong_an)

                continue

            if L1max < 23:

                option_2_1_1_2_cot(doc, phuong_an)

                continue

            option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot_trong_cac_vung_font13(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 11:

                option_1_2_cot(doc, phuong_an)

                continue

            if L1max < 21:

                option_2_1_1_2_cot(doc, phuong_an)

                continue

            option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot_trong_cac_vung_font14(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 10:

                option_1_2_cot(doc, phuong_an)

                continue

            if L1max < 19:

                option_2_1_1_2_cot(doc, phuong_an)

                continue

            option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot_trong_cac_vung_Japan(doc, text_start, text_end):

    all_indices_cau = tim_cau_trong_phan(doc, text_start, text_end)

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

            if not len(phuong_an) > 2:

                continue

            phuong_an = phuong_an[:-1]

            L1max = max_len_option(doc, phuong_an)

            if L1max < 5:

                option_1_2_cot(doc, phuong_an)

                continue

            if L1max < 9:

                option_2_1_1_2_cot(doc, phuong_an)

                continue

            option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_2_cot_trong_cac_vung_font13(doc, 'S1@', 'E1@')

        return None

    if size_font == 14:

        sap_xep_phan_I_2_cot_trong_cac_vung_font14(doc, 'S1@', 'E1@')

        return None

    sap_xep_phan_I_2_cot_trong_cac_vung(doc, 'S1@', 'E1@')





def sap_xep_phan_I_2_cot_EN(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_2_cot_trong_cac_vung_font13(doc, '<S@>', '<E@>')

        return None

    if size_font == 14:

        sap_xep_phan_I_2_cot_trong_cac_vung_font14(doc, '<S@>', '<E@>')

        return None

    sap_xep_phan_I_2_cot_trong_cac_vung(doc, '<S@>', '<E@>')





def sap_xep_phan_I_2_cot_Japan(doc):

    sap_xep_phan_I_2_cot_trong_cac_vung_Japan(doc, '<S@>', '<E@>')





def lam_dep_phan_II(doc):

    all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, 'S2@', 'E2@')

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_II(doc, cau[k], cau[k + 1])

            phuong_an = phuong_an[:-1]

            for p in phuong_an:

                insert_tab_before_first_child(doc, p)

                tab_btp_4_at(doc, p)





def lam_dep_phan_IV(doc):

    all_indices_cau = mdocx_f.tim_cau_trong_phan(doc, 'S4@', 'E4@')

    for idx in range(len(all_indices_cau) - 1, -1, -1):

        cau = all_indices_cau[idx]

        for k in range(len(cau) - 2, -1, -1):

            phuong_an = mdocx_f.tim_phuong_an_trong_cau_phan_II(doc, cau[k], cau[k + 1])

            phuong_an = phuong_an[:-1]

            for p in phuong_an:

                insert_tab_before_first_child(doc, p)

                tab_btp_4_at(doc, p)





def Xoa_dong_chua_tu(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not text in paragraphs[p].text:

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def Xoa_dong_start_with(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not paragraphs[p].text.startswith(text):

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def replace_text_in_doc(doc, old_text, new_text):

    '''Thay thế old_text bằng new_text trong toàn bộ tài liệu, giữ nguyên định dạng'''

    replace_in_paragraphs(doc.paragraphs, old_text, new_text)

    replace_in_tables(doc.tables, old_text, new_text)

    for section in doc.sections:

        replace_in_paragraphs(section.header.paragraphs, old_text, new_text)

        replace_in_tables(section.header.tables, old_text, new_text)

        replace_in_paragraphs(section.footer.paragraphs, old_text, new_text)

        replace_in_tables(section.footer.tables, old_text, new_text)





def replace_in_paragraphs(paragraphs, old_text, new_text):

    '''Thay thế văn bản trong danh sách đoạn văn'''

    for para in paragraphs:

        for run in para.runs:

            if not old_text in run.text:

                continue

            run.text = run.text.replace(old_text, new_text)





def replace_in_tables(tables, old_text, new_text):

    '''Thay thế văn bản trong tất cả các bảng'''

    for table in tables:

        for row in table.rows:

            for cell in row.cells:

                replace_in_paragraphs(cell.paragraphs, old_text, new_text)

                replace_in_tables(cell.tables, old_text, new_text)





def update_made(doc, new_text):

    """Thay thế tất cả '<made>' trong tài liệu"""

    replace_text_in_doc(doc, '<made>', new_text)





def lay_so_trang_tu_footer(doc):

    last_section = doc.sections[0]

    footer = last_section.footer

    num_pages = None

    for paragraph in footer.paragraphs:

        print(paragraph.text)





def get_page_count_from_docx(docx_path):

    pass

# WARNING: Decompyle incomplete





def update_so_trang(docx_path):

    num_pages = get_page_count_from_docx(docx_path)

    new_text = f'''{num_pages:02}'''

    doc = Document(docx_path)

    replace_text_in_doc(doc, '<sotrang>', new_text)

    doc.save(docx_path)





def canh_before_after_for_cau(doc):

    pattern1 = '^Question [0-9]{1,}[.:]'

    pattern2 = '^Câu [0-9]{1,}[.:]'

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    indices_cau = []

    for j in range(len(paragraphs)):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not re.match(pattern1, text) and re.match(pattern2, text):

            continue

        indices_cau.append(j)

    for p in indices_cau:

        canh_before_after(doc, p)





def canh_before_after_all(doc):

    for p in range(len(doc.paragraphs)):

        paragraph = doc.paragraphs[p]

        paragraph.paragraph_format.space_before = Pt(0)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def save_and_close_docx(doc, output_file):

    doc.save(output_file)





def chuanhoadocx(doc, output_file, new_text, page_kind, bo_loi_giai = ('A4', 'YES')):

    Danh_STT_Cau_All(doc)

    STT_ABCD(doc)

    STT_abcd_nho(doc)

    STT_thang(doc)

    if bo_loi_giai == 'YES':

        docxtool.xoa_red_dap_an_docx(doc)

    if page_kind == '2C':

        sap_xep_phan_I_2_cot(doc)

        page_2_cot_Mix(doc)

        resize_images_to_column_width(doc)

    else:

        sap_xep_phan_I(doc)

    lam_dep_phan_II(doc)

    lam_dep_phan_IV(doc)

    update_made(doc, new_text)

    if bo_loi_giai == 'YES':

        docxtool.xoa_loi_giai_lay_de(doc)

        Xoa_dong_chua_tu(doc, '@')

        Xoa_dong_start_with(doc, 'ĐS:')

    else:

        docxtool.xoa_gach_chan_tab_trong_docx(doc)

    save_and_close_docx(doc, output_file)

    return None

# WARNING: Decompyle incomplete





def chuanhoadocx_EN(doc, output_file, new_text, page_kind, bo_loi_giai = ('A4', 'YES')):

    Danh_STT_Cau_All_EN(doc)

    STT_ABCD_EN(doc)

    STT_thang(doc)

    if bo_loi_giai == 'YES':

        docxtool.xoa_red_dap_an_docx(doc)

    if page_kind == '2C':

        sap_xep_phan_I_2_cot_EN(doc)

        page_2_cot_Mix(doc)

        resize_images_to_column_width(doc)

    else:

        sap_xep_phan_I_EN(doc)

    if bo_loi_giai == 'YES':

        docxtool.xoa_loi_giai_lay_de(doc)

        Xoa_dong_chua_tu(doc, '@')

    else:

        docxtool.xoa_gach_chan_tab_trong_docx(doc)

    update_made(doc, new_text)

    save_and_close_docx(doc, output_file)

    return None

# WARNING: Decompyle incomplete





def chuanhoadocx_Japan(doc, output_file, new_text, page_kind, bo_loi_giai = ('A4', 'YES')):

    Danh_STT_Cau_All_EN(doc)

    STT_ABCD_EN(doc)

    STT_thang(doc)

    if bo_loi_giai == 'YES':

        docxtool.xoa_red_dap_an_docx(doc)

    if page_kind == '2C':

        sap_xep_phan_I_2_cot_Japan(doc)

        page_2_cot_Mix(doc)

        resize_images_to_column_width(doc)

    else:

        sap_xep_phan_I_Japan(doc)

    if bo_loi_giai == 'YES':

        docxtool.xoa_loi_giai_lay_de(doc)

        Xoa_dong_chua_tu(doc, '@')

    else:

        docxtool.xoa_gach_chan_tab_trong_docx(doc)

    update_made(doc, new_text)

    save_and_close_docx(doc, output_file)

    return None

# WARNING: Decompyle incomplete





def chuanhoadocx_ngoai_ngu(doc, output_file, new_text, page_kind, bo_loi_giai, langue = ('A4', 'YES', 'english')):

    if langue == 'japan':

        chuanhoadocx_Japan(doc, output_file, new_text, page_kind, bo_loi_giai)

        return None

    chuanhoadocx_EN(doc, output_file, new_text, page_kind, bo_loi_giai)



