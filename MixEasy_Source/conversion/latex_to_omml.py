import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os
import re
from latex2mathml.converter import convert as latex2mathml
import mathml2omml
from docx.oxml import parse_xml
from lxml.etree import XMLSyntaxError

def _braces_balanced(s = None):
    c = 0
    for ch in s:
        if ch == '{':
            c += 1
            continue
        if not ch == '}':
            continue
        c -= 1
        if not c < 0:
            continue
        s
        return False
    return c == 0


def preprocess_latex(latex = None):
    s = latex.strip()
    if s.startswith('{') and s.endswith('}') and _braces_balanced(s):
        s = s[1:-1].strip()
    s = re.sub('\\\\bar\\{([^}]*)\\}', '\\\\overline{\\1}', s)
    s = re.sub('\\\\bar([A-Za-z0-9])', '\\\\overline{\\1}', s)
    s = s.replace('\\dfrac', '\\frac')
    s = s.replace('\\tilde', '\\widetilde')
    s = s.replace('\\hat', '\\widehat')
    return s


def insert_equation(paragraph, latex_src, fail_log = (None,)):
    try_variants = [
        latex_src.strip(),
        preprocess_latex(latex_src)]
    last_err = 'No valid variant to try'
    for idx, latex_try in enumerate(try_variants):
        if not latex_try:
            continue
        mathml = latex2mathml(latex_try)
        omml = mathml2omml.convert(mathml)
        xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"\n             xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</w:r>'''
        r = parse_xml(xml)
        paragraph._element.append(r)
        enumerate(try_variants)
        return True
# WARNING: Decompyle incomplete


def replace_double_dollar(text = None):
    '''Thay táº¥t cáº£ $$ thÃ nh $ (khÃ´ng Ä‘á»¥ng ná»™i dung bÃªn trong).'''
    return text.replace('$$', '$')


def merge_block_latex(doc):
    i = 0
    paras = doc.paragraphs
    if i < len(paras):
        p = paras[i]
        for r in p.runs:
            if not '$$' in r.text:
                continue
            r.text = replace_double_dollar(r.text)
        i += 1
        if i < len(paras):
            continue
    i = 0
    if i < len(paras):
        p = paras[i]
        if p.text.strip() == '$':
            parts = []
            j = i + 1
            if j < len(paras) and paras[j].text.strip() != '$':
                parts.append(paras[j].text.strip())
                j += 1
                if j < len(paras) and paras[j].text.strip() != '$':
                    continue
            if j < len(paras) and paras[j].text.strip() == '$':
                combined = f'''${' '.join(parts)}$'''
                for r in p.runs:
                    r.text = ''
                p.add_run(combined)
                for k in range(i + 1, j + 1):
                    for r in paras[k].runs:
                        r.text = ''
                i = j
        i += 1
        if i < len(paras):
            continue
        return None


def process_file(filepath):
    doc = Document(filepath)
    file_failed = []
    merge_block_latex(doc)
    for p in doc.paragraphs:
        text = p.text
        if not '$' in text and '\\[' in text:
            continue
        pattern = '(\\$.*?\\$|\\\\\\[.*?\\\\\\])'
        parts = re.split(pattern, text)
        for run in p.runs:
            run.clear()
        p.clear()
        for part in parts:
            if not part:
                continue
            if part.startswith('$') and part.endswith('$'):
                latex_expr = part[1:-1]
                ok = insert_equation(p, latex_expr, fail_log = file_failed)
                if ok:
                    continue
                p.add_run(part)
                continue
            if part.startswith('\\[') and part.endswith('\\]'):
                latex_expr = part[2:-2]
                ok = insert_equation(p, latex_expr, fail_log = file_failed)
                if ok:
                    continue
                p.add_run(part)
                continue
            p.add_run(part)
    newpath = os.path.splitext(filepath)[0] + '_equation.docx'
    doc.save(newpath)
    if file_failed:
        return {
            'file': filepath,
            'failed': file_failed }
# WARNING: Decompyle incomplete


def browse_files(listbox):
    files = filedialog.askopenfilenames(filetypes = [
        ('Word files', '*.docx')])
    for f in files:
        listbox.insert(tk.END, f)


def clear_list(listbox):
    listbox.delete(0, tk.END)


def convert_files(listbox):
    files = listbox.get(0, tk.END)
    if not files:
        messagebox.showwarning('No files', 'Please select at least one file.')
        return None
    failed_overall = []
    for f in files:
        result = process_file(f)
        if not result:
            continue
        failed_overall.append(result)
    if not failed_overall:
        messagebox.showinfo('Done', 'All files converted successfully!')
        return None
    msg = 'Some files had errors:\n\n'
    for res in failed_overall:
        msg += f'''{res['file']}:\n'''
        for latex_src, err in res['failed']:
            msg += f'''  - {latex_src} -> {err}\n'''
    messagebox.showwarning('Lá»—i, kiá»ƒm tra xem Ä‘Ã£ táº¯t word chÆ°a', msg)


def gui_convert_latex_to_OMML(root):
    pass
# WARNING: Decompyle incomplete

if __name__ == '__main__':
    root = tk.Tk()
    gui_convert_latex_to_OMML(root)
    root.mainloop()
    return None
