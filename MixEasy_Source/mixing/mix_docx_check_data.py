import os

import tkinter as tk

from tkinter import ttk, Label, filedialog, Entry, Button, Frame, Listbox, Scrollbar, messagebox

from tkinter import simpledialog

from docx import Document

from docx.text.paragraph import Paragraph

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import RGBColor

from docx.shared import Inches, Pt

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from lxml import etree



ElementTree

from xml.etree.ElementTree import QName

import xml.etree.ElementTree, etree

import re

import zipfile

import bisect

from copy import deepcopy



def find_phan(doc, text_start, text_end):

    '''Tìm chỉ mục của các đoạn văn bắt đầu bằng text_start và text_end'''

    indices_S = []

    indices_E = []

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, para in enumerate(paragraphs):

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if text.startswith(text_start):

            indices_S.append(i)

            continue

        if not text.startswith(text_end):

            continue

        indices_E.append(i)

    return (indices_S, indices_E)





def tim_cau_trong_phan(doc, text_start, text_end):

    '''Tìm danh sách các câu hỏi trong khoảng từ text_start đến text_end'''

    (indices_S, indices_E) = find_phan(doc, text_start, text_end)

    all_indices_cau = []

    if len(indices_S) != len(indices_E):

        return None

    if len(indices_S) > 0:

        pattern1 = '^Câu [0-9]{1,}[.:]'

        pattern2 = '^Question [0-9]{1,}[.:]'

        body = doc.element.body

        paragraphs = list(body.iterchildren())

        for start, end in zip(indices_S, indices_E):

            indices_cau = []

            for j in range(start, end):

                para = paragraphs[j]

                if not para.tag.endswith('p'):

                    continue

                paragraph = Paragraph(para, doc)

                text = paragraph.text

                if not re.match(pattern1, text) and re.match(pattern2, text):

                    continue

                indices_cau.append(j)

            indices_cau.append(end)

            all_indices_cau.append(indices_cau)

    return all_indices_cau





def tim_phuong_an_trong_cau_phan_I(doc, from_a, to_b):

    """Tìm phương án trong câu hỏi theo định dạng 'A.', 'B.', 'C.', 'D.' từ from_a đến to_b"""

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    to_end = to_b

    for j in range(from_a, to_b):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text.strip()

        if not text == 'Lời giải':

            continue

        to_end = j

        range(from_a, to_b)

    labels = [

        'A.',

        'B.',

        'C.',

        'D.']

    indices_phuong_an = []

    for j in range(from_a, min(to_end, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def tim_phuong_an_dung_trong_cau_phan_I(doc, from_a, to_b):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    to_end = to_b

    for j in range(from_a, to_b):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text.strip()

        if not text == 'Lời giải':

            continue

        to_end = j

        range(from_a, to_b)

    labels = [

        'A.',

        'B.',

        'C.',

        'D.']

    indices_phuong_an = []

    for j in range(from_a, min(to_end, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith(tuple(labels)):

            continue

        runs = paragraph.runs

        if not runs[0].font.underline:

            runs[0].font.underline

            if runs[0].font.color:

                runs[0].font.color

        Dap_an = runs[0].font.color.rgb == RGBColor(255, 0, 0)

        if not Dap_an:

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def tim_phuong_an_trong_cau_phan_II(doc, from_a, to_b):

    """Tìm phương án trong câu hỏi theo định dạng 'a)', 'b)', 'c)', 'd)' từ from_a đến to_b"""

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    to_end = to_b

    for j in range(from_a, to_b):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text.strip()

        if not text == 'Lời giải':

            continue

        to_end = j

        range(from_a, to_b)

    labels = [

        'a)',

        'b)',

        'c)',

        'd)']

    indices_phuong_an = []

    for j in range(from_a, min(to_end, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def tim_dap_so_trong_cau_phan_III(doc, from_a, to_b):

    """Tìm phương án trong câu hỏi theo định dạng 'a)', 'b)', 'c)', 'd)' từ from_a đến to_b"""

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    indices_phuong_an = []

    for j in range(from_a, min(to_b, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith('ĐS:'):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def Check_math_dap_so(doc, p, messages, nhom, cau):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    if p < 0 or p >= len(paragraphs):

        raise IndexError('Chỉ số p không hợp lệ.')

    para = paragraphs[p]

    if not para.tag.endswith('p'):

        return 0

    nsmap = doc.element.nsmap.copy()

    if 'a' not in nsmap:

        nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    if 'm' not in nsmap:

        nsmap['m'] = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    if 'o' not in nsmap:

        nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

    math_objects = para.findall('.//m:oMathPara', namespaces = nsmap)

    if not math_objects:

        math_objects = para.findall('.//m:oMath', namespaces = nsmap)

    if math_objects:

        messages.append(f'''Phần III, Nhóm {nhom + 1}, Câu {cau + 1}: Đáp án được đánh bằng Equation''')

# WARNING: Decompyle incomplete





def check_file_mau(doc, mode = ('VN',)):

    pass

# WARNING: Decompyle incomplete





def check_ki_hieu(doc, messages, text_start, text_end):

    indices_S = []

    indices_E = []

    for i, paragraph in enumerate(doc.paragraphs, start = 1):

        text = paragraph.text

        if text.startswith(text_start):

            indices_S.append(i)

            continue

        if not text.startswith(text_end):

            continue

        indices_E.append(i)

    if len(indices_S) != len(indices_E):

        messages.append(f'''Số kí hiệu {text_start} và {text_end} không bằng nhau- KHÔNG hợp lệ''')

        return None





def check_du_lieu_sau_xuongdong_P1(doc, messages):

    all_indices_cau = tim_cau_trong_phan(doc, 'S1@', 'E1@')

    for i, indices_cau in enumerate(all_indices_cau):

        for k in range(len(indices_cau) - 1):

            phuong_an = tim_phuong_an_trong_cau_phan_I(doc, indices_cau[k], indices_cau[k + 1])

            if len(phuong_an) != 4:

                messages.append(f'''Phần I, Nhóm {i + 1}, Câu {k + 1} có số phương án là {len(phuong_an)}''')

                messages.append('Nếu có đáp án bị nằm trong bảng thì số lượng báo không chính xác và sẽ ảnh hưởng đến nhiều việc khác')

            phuong_an_dung = tim_phuong_an_dung_trong_cau_phan_I(doc, indices_cau[k], indices_cau[k + 1])

            if not len(phuong_an_dung) != 1:

                continue

            messages.append(f'''Phần I, Nhóm {i + 1}, Câu {k + 1} có số phương án đúng là {len(phuong_an_dung)}''')





def check_du_lieu_sau_xuongdong_P1_EN(doc, messages):

    all_indices_cau = tim_cau_trong_phan(doc, '<S@>', '<E@>')

    for i, indices_cau in enumerate(all_indices_cau):

        for k in range(len(indices_cau) - 1):

            phuong_an = tim_phuong_an_trong_cau_phan_I(doc, indices_cau[k], indices_cau[k + 1])

            if len(phuong_an) != 4:

                messages.append(f'''Phần I, Nhóm {i + 1}, Câu {k + 1} có số phương án là {len(phuong_an)}''')

                messages.append('Nếu có đáp án bị nằm trong bảng thì số lượng báo không chính xác và sẽ ảnh hưởng đến nhiều việc khác')

            phuong_an_dung = tim_phuong_an_dung_trong_cau_phan_I(doc, indices_cau[k], indices_cau[k + 1])

            if not len(phuong_an_dung) != 1:

                continue

            messages.append(f'''Phần I, Nhóm {i + 1}, Câu {k + 1} có số phương án đúng là {len(phuong_an_dung)}''')





def check_du_lieu_sau_xuongdong_P2(doc, messages):

    all_indices_cau = tim_cau_trong_phan(doc, 'S2@', 'E2@')

    for i, indices_cau in enumerate(all_indices_cau):

        for k in range(len(indices_cau) - 1):

            phuong_an = tim_phuong_an_trong_cau_phan_II(doc, indices_cau[k], indices_cau[k + 1])

            if not len(phuong_an) != 4:

                continue

            messages.append(f'''Phần II, Nhóm {i + 1}, Câu {k + 1} có số phương án là {len(phuong_an)}''')

            messages.append('Nếu có đáp án bị nằm trong bảng thì số lượng báo không chính xác và sẽ ảnh hưởng đến nhiều việc khác')





def check_du_lieu_sau_xuongdong_P3(doc, messages):

    all_indices_cau = tim_cau_trong_phan(doc, 'S3@', 'E3@')

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, indices_cau in enumerate(all_indices_cau):

        for k in range(len(indices_cau) - 1):

            indices_dap_so = []

            for j in range(indices_cau[k], indices_cau[k + 1]):

                para = paragraphs[j]

                if not para.tag.endswith('p'):

                    continue

                paragraph = Paragraph(para, doc)

                text = paragraph.text

                if not text.startswith('ĐS:'):

                    continue

                indices_dap_so.append(j)

                Check_math_dap_so(doc, j, messages, i, k)

            if not len(indices_dap_so) != 1:

                continue

            messages.append(f'''Phần III, Nhóm {i + 1}, Câu {k + 1} có số đáp án là {len(indices_dap_so)}''')





def check_sentences_in_tables_docx(doc, messages):

    pattern = re.compile('^Câu \\d+[.:]')

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    text = para.text.strip()

                    if not pattern.match(text):

                        continue

                    messages.append("Có Câu nằm trong bảng, nếu nó là 1 câu của đề thi thì phải đưa nó ra ngoài bảng đó nhé! \n Có thể mở file lên, seclect vùng đó và dùng 'Convert Table to text (F1)' trong tool 1 ")

                    cell.paragraphs





def check_data_docx_mix(doc):

    messages = []

    check_file_mau(doc, mode = 'VN')

    check_ki_hieu(doc, messages, 'S1@', 'E1@')

    check_ki_hieu(doc, messages, 'S2@', 'E2@')

    check_ki_hieu(doc, messages, 'S3@', 'E3@')

    check_ki_hieu(doc, messages, 'S4@', 'E4@')

    check_ki_hieu(doc, messages, '<SCHUM@>', '<ECHUM@>')

    check_ki_hieu(doc, messages, '<SCHUM@><CHON>', '<ECHUM@><CHON>')

    check_du_lieu_sau_xuongdong_P1(doc, messages)

    check_du_lieu_sau_xuongdong_P2(doc, messages)

    check_du_lieu_sau_xuongdong_P3(doc, messages)

    check_sentences_in_tables_docx(doc, messages)

    if messages:

        messages.append('SAU KHI SỬA CHỮA HÃY KIỂM TRA LẠI')

        return (False, messages)

    return (None, [])

# WARNING: Decompyle incomplete





def check_data_docx_mix_EN(doc):

    messages = []

    check_file_mau(doc, mode = 'EN')

    check_ki_hieu(doc, messages, '<S@>', '<E@>')

    check_ki_hieu(doc, messages, '<SNHOM@>', '<ENHOM@>')

    check_ki_hieu(doc, messages, '<SNHOM@><CD>', '<ENHOM@><CD>')

    check_ki_hieu(doc, messages, '<SNHOM@><DC>', '<ENHOM@><DC>')

    check_ki_hieu(doc, messages, '<SNHOM@><C>', '<ENHOM@><C>')

    check_ki_hieu(doc, messages, '<SNHOM@><D>', '<ENHOM@><D>')

    check_du_lieu_sau_xuongdong_P1_EN(doc, messages)

    check_sentences_in_tables_docx(doc, messages)

    if messages:

        messages.append('SAU KHI SỬA CHỮA HÃY KIỂM TRA LẠI')

        return (False, messages)

    return (None, [])

# WARNING: Decompyle incomplete





def browse_file(entry_file_path):

    input_file = filedialog.askopenfilename(filetypes = [

        ('Word files', '*.docx')])

    if not input_file:

        return None

    input_file = os.path.normpath(input_file)

    entry_file_path.delete(0, tk.END)

    entry_file_path.insert(tk.END, input_file)





def gui_check_du_lieu_sau_xuongdong2():

    pass

# WARNING: Decompyle incomplete





def gui_check_du_lieu_sau_xuongdong():

    pass

# WARNING: Decompyle incomplete





def gui_check_du_lieu_sau_xuongdong_EN():

    pass

# WARNING: Decompyle incomplete



if __name__ == '__main__':

    root = tk.Tk()

    gui_check_du_lieu_sau_xuongdong()

    return None

