import os

import tkinter as tk

from tkinter import ttk, Label, filedialog, Entry, Button, Frame, Listbox, Scrollbar, messagebox

from tkinter import simpledialog

from docx import Document

from docx.text.paragraph import Paragraph

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import RGBColor

from docx.shared import Inches, Pt

from docx.oxml import parse_xml

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from docx.text.run import Run

from lxml import etree



ElementTree

from xml.etree.ElementTree import QName

import xml.etree.ElementTree, etree

import re

import zipfile

import bisect

from copy import deepcopy

import math



def InchesToPoints(inches):

    return Pt(inches * 72)





def LinesToPoints(lines):

    return Pt(lines * 12)





def open_doc_off_python_docx(doc_path):

    doc = Document(doc_path)

    return doc





def remove_last_element(lst):

    return lst[:-1]





def is_paragraph_numbered(paragraph):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    indices = []

# WARNING: Decompyle incomplete





def has_numbered_paragraph(doc):

    '''

    Trả về True nếu tài liệu có ít nhất một đoạn văn <w:p> được đánh số tự động.

    Ngược lại, trả về False.

    '''

    body = doc.element.body

    paragraphs = list(body.iterchildren())

# WARNING: Decompyle incomplete





def has_numbered_paragraph_Mix(doc, messages):

    '''

    Trả về True nếu tài liệu có ít nhất một đoạn văn <w:p> được đánh số tự động.

    Ngược lại, trả về False.

    '''

    body = doc.element.body

    paragraphs = list(body.iterchildren())

# WARNING: Decompyle incomplete





def xoa_ky_tu_an_docx(doc):

    invisible_map = {

        '​': '',

        '‬': '',

        '⁠': '',

        '﻿': '',

        '‌': '',

        '‍': '',

        '\xc2\xa0': ' ' }

    for para in doc.paragraphs:

        for run in para.runs:

            text = run.text

            for char, replacement in invisible_map.items():

                if not char in text:

                    continue

                run.text = run.text.replace(char, replacement)





def thay_ky_tu_gia_tao_ABCD_docx(doc):

    invisible_map = {

        'Α': 'A',

        'Β': 'B',

        'А': 'A',

        'В': 'B',

        'С': 'C',

        'а': 'a',

        'в': 'b',

        'с': 'c' }

    for para in doc.paragraphs:

        for run in para.runs:

            text = run.text

            for char, replacement in invisible_map.items():

                if not char in text:

                    continue

                run.text = run.text.replace(char, replacement)





def xoa_dong_trang_docx_new(doc):

    pass

# WARNING: Decompyle incomplete





def Xoa_dong_start_with_tool(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not paragraphs[p].text.startswith(text):

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def xoa_red_dap_an_docx_tool(doc):

    labels = [

        'A',

        'B',

        'C',

        'D',

        'a',

        'b',

        'c',

        'd']

    for para in doc.paragraphs:

        for run in para.runs:

            text = run.text.strip()

            if not text.strip().startswith(tuple(labels)):

                continue

            if not run.font.color.rgb == RGBColor(255, 0, 0) and run.font.underline:

                continue

            run.font.underline = None

            run.font.color.rgb = RGBColor(0, 0, 255)

    Xoa_dong_start_with_tool(doc, 'ĐS:')





def xoa_red_dap_an_docx(doc):

    labelf = [

        'A',

        'B',

        'C',

        'D',

        'a',

        'b',

        'c',

        'd']

    labels = [

        'A.',

        'B.',

        'C.',

        'D.',

        'a)',

        'b)',

        'c)',

        'd)']

    for para in doc.paragraphs:

        text = para.text.strip()

        if not text.startswith(tuple(labels)):

            continue

        run = para.runs[0]

        run_text = run.text.strip()

        if not run_text.startswith(tuple(labelf)):

            continue

        run.font.underline = None

        run.font.color.rgb = RGBColor(0, 0, 255)





def xoa_red_dap_an_docx_sau(doc):

    labelf = [

        'A',

        'B',

        'C',

        'D',

        'a',

        'b',

        'c',

        'd']

    labels = [

        'A.',

        'B.',

        'C.',

        'D.',

        'a)',

        'b)',

        'c)',

        'd)']

    for para in doc.paragraphs:

        text = para.text.strip()

        if not text.startswith(tuple(labels)):

            continue

        for run in para.runs:

            run_text = run.text.strip()

            if not run_text.startswith(tuple(labelf)):

                continue

            if not run.font.color.rgb == RGBColor(255, 0, 0):

                continue

            run.font.underline = None

            run.font.color.rgb = RGBColor(0, 0, 255)





def xoa_dong_trang_docx(doc):

    paragraph = doc.paragraphs

    index2 = len(paragraph)

    for i in range(index2 - 1, 0, -1):

        para = paragraph[i]

        if para.text.strip():

            continue

        if para.runs:

            continue

        p = para._element

        p.getparent().remove(p)





def remove_paragraphs_starting_with(doc, prefix):

    paragraphs_to_remove = []

    for para in doc.paragraphs:

        if not para.text.startswith(prefix):

            continue

        paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:

        p = para._element

        p.getparent().remove(p)





def xoa_QC(doc_path, page_kind = ('A4',)):

    if not os.path.exists(doc_path):

        print(f'''File không tồn tại: {doc_path}''')

        return None

    doc = Document(doc_path)

    remove_paragraphs_starting_with(doc, 'Evaluation Warning')

    if page_kind == '2C':

        set_two_columns_with_spacing(doc, spacing_in_inches = 0.1)

    doc.save(doc_path)

    return None

# WARNING: Decompyle incomplete





def xoa_loi_giai_acong_dapso_docx_for_mix_spr(output_file_loi_giai, output_file):

    doc = Document(output_file_loi_giai)

    xoa_loi_giai_lay_de(doc)

    xoa_red_dap_an_docx_sau(doc)

    Xoa_dong_chua_tu(doc, '@')

    Xoa_dong_start_with(doc, 'ĐS:')

    doc.save(output_file)





def xoa_loi_giai_acong_for_mix_docx_EN(output_file_loi_giai, output_file):

    doc = Document(output_file_loi_giai)

    xoa_loi_giai_lay_de(doc)

    xoa_red_dap_an_docx_sau(doc)

    Xoa_dong_chua_tu(doc, '@')

    doc.save(output_file)





def xoa_QC_loi_giai_acong_dapso(doc_path, page_kind = ('A4',)):

    if not os.path.exists(doc_path):

        print(f'''File không tồn tại: {doc_path}''')

        return None

    doc = Document(doc_path)

    remove_paragraphs_starting_with(doc, 'Evaluation Warning')

    xoa_loi_giai_lay_de(doc)

    xoa_red_dap_an_docx_sau(doc)

    Xoa_dong_chua_tu(doc, '@')

    Xoa_dong_start_with(doc, 'ĐS:')

    if page_kind == '2C':

        set_two_columns_with_spacing(doc, spacing_in_inches = 0.1)

    doc.save(doc_path)

    return None

# WARNING: Decompyle incomplete





def set_two_columns_with_spacing(doc, spacing_in_inches):

    sectPr = doc.sections[0]._sectPr

    cols = OxmlElement('w:cols')

    cols.set(qn('w:num'), '2')

    cols.set(qn('w:sep'), '1')

    spacing_in_twips = int(spacing_in_inches * 1440)

    cols.set(qn('w:space'), str(spacing_in_twips))

    sectPr.append(cols)





def tim_body_HET(doc):

    pattern1 = 'HẾT'

    pattern2 = 'Hết'

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for j in range(len(body)):

        para = paragraphs[j]

        text = (lambda .0: pass# WARNING: Decompyle incomplete

)(para.iter()()).strip()

        if not re.search(pattern1, text) and re.search(pattern2, text):

            continue

        

        return ''.join, j

    return 0





def xoa_huong_dan_lay_de_docx(doc, messages):

    '''Trộn ngẫu nhiên các nhóm đoạn văn dựa trên danh sách indices'''

    start = tim_body_HET(doc)

    if start > 0:

        body = doc.element.body

        paragraphs = list(body.iterchildren())

        for i in range(start + 1, len(body) - 1):

            body.remove(paragraphs[i])

        xoa_dong_trang_docx_new(doc)

        return None

    messages.append('Không có từ HẾT ngăn cách 2 phần')





def xoa_de_lay_huong_dan_docx(doc, messages):

    end = tim_body_HET(doc)

    if end > 0:

        body = doc.element.body

        paragraphs = list(body.iterchildren())

        for i in range(0, end + 1):

            body.remove(paragraphs[i])

        xoa_dong_trang_docx_new(doc)

        return None

    messages.append('Không có từ HẾT ngăn cách 2 phần')





def tim_cau_body_all(doc, selected_headings):

    pass

# WARNING: Decompyle incomplete





def chuanhoa_loigiai_docx(doc):

    pattern = re.compile('\\s*(giải|lời giải|hướng dẫn|hướng dẫn giải)[ \\t]*[.:]*[ \\t]*', re.IGNORECASE)

    for para in doc.paragraphs:

        text = para.text.strip()

        if not pattern.fullmatch(text):

            continue

        para.clear()

        run = para.add_run('Lời giải')

        run.bold = True

        run.font.color.rgb = RGBColor(0, 0, 255)

        run.font.name = 'Times New Roman'

        run.font.size = Pt(12)

        para.alignment = 1





def tim_loi_giai_body_all(doc):

    """Tìm tất cả vị trí chứa 'Question X.' hoặc 'Câu X.' trong tài liệu"""

    pattern = re.compile('\\s*(giải|lời giải|hướng dẫn|hướng dẫn giải)[ \\t]*[.:]*[ \\t]*', re.IGNORECASE)

    indices_LG = []

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, para in enumerate(paragraphs):

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text.strip().lower()

        if not pattern.fullmatch(text):

            continue

        indices_LG.append(i)

    return indices_LG





def tim_chi_so_loi_giai_para(doc):

    """Tìm tất cả vị trí chứa 'Question X.' hoặc 'Câu X.' trong tài liệu"""

    pattern = re.compile('\\s*(giải|lời giải|hướng dẫn|hướng dẫn giải)[ \\t]*[.:]*[ \\t]*', re.IGNORECASE)

    indices_LG = []

    for i, para in enumerate(doc.paragraphs):

        text = para.text.strip().lower()

        if not pattern.fullmatch(text):

            continue

        indices_LG.append(i)

    return indices_LG





def tao_danh_sach_loigiai_cau(a, b):

    pass

# WARNING: Decompyle incomplete





def xoa_loi_giai_lay_de(doc, selected_headings = ([],)):

    indices = tim_chi_so_cau_phan(doc)

    delete_run0_empty_indices_ok(doc, indices)

    indices = tim_chi_so_loi_giai_para(doc)

    delete_run0_empty_indices_ok(doc, indices)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    a = tim_cau_body_all(doc, selected_headings)

    b = tim_loi_giai_body_all(doc)

    c = tao_danh_sach_loigiai_cau(a, b)

    if len(b) == len(c):

        for start, end in reversed(list(zip(b, c))):

            for i in range(end - 1, start - 1, -1):

                if paragraphs[i].getparent() == body:

                    body.remove(paragraphs[i])

        continue

    xoa_dong_trang_docx_new(doc)

    return None

# WARNING: Decompyle incomplete





def xoa_header(doc):

    

    def clear_all_content(part):

        for child in list(part._element):

            part._element.remove(child)



    for section in doc.sections:

        headers = [

            section.header,

            section.first_page_header,

            section.even_page_header]

        footers = [

            section.footer,

            section.first_page_footer,

            section.even_page_footer]

        for header in headers:

            clear_all_content(header)

        for footer in footers:

            clear_all_content(footer)





def xoa_tac_gia(doc):

    pass

# WARNING: Decompyle incomplete





def xoa_Chu_thich(doc):

    pattern = re.compile('\\s*((?:câu|question|ví dụ)\\s*\\d+[ \\t]*[.:][ \\t]*)([\\[\\(].*?[\\]\\)])', re.IGNORECASE)

    indices_phan = tim_chi_so_phan(doc)

    if len(indices_phan) > 1:

        for k in range(len(indices_phan) - 1):

            indices_cau = tim_chi_so_cau_from_a_to_b(doc, indices_phan[k], indices_phan[k + 1])

            STT = 1

            for i in indices_cau:

                para = doc.paragraphs[i]

                runs = para.runs

                full_text = ''

                match_run_indices = []

                for idx, run in enumerate(runs):

                    if not run.text:

                        enumerate(runs)

                    else:

                        full_text += run.text

                        match_run_indices.append(idx)

                        if not re.match(pattern, full_text):

                            continue

                        indices_cau

            match = re.match(pattern, full_text)

            if not match:

                continue

            matched_text = match.group()

            tu_khoa = match.group(1).capitalize()

            remaining_text = full_text[len(matched_text):]

            last_run_idx = match_run_indices[-1]

            last_run = runs[last_run_idx]

            cloned_run = deepcopy(last_run)

            runs[0].text = f'''{tu_khoa}'''

            for idx in reversed(match_run_indices[1:]):

                para._element.remove(runs[idx]._element)

            if remaining_text:

                cloned_run.text = remaining_text

                runs[0]._element.addnext(cloned_run._element)

            STT += 1

        continue

        return None





def move_table_paragraphs_to_top2(doc):

    body = doc.element.body

    tables = list(body.iterchildren(tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'))

    if not tables:

        print('Không có bảng nào.')

        return None

    table = doc.tables[0]

    table_element = tables[0]

    para_texts = []

    for row in table.rows:

        for cell in row.cells:

            for para in cell.paragraphs:

                text = para.text.strip()

                if not text:

                    continue

                para_texts.append(text)

    table_element.getparent().remove(table_element)

    for text in reversed(para_texts):

        new_p = doc.add_paragraph()

        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = new_p.add_run(text)

        run.bold = True

        run.font.color.rgb = RGBColor(0, 0, 255)

        body.insert(0, new_p._element)





def xoa_bang_dau_giu_lai_tieu_de(doc):

    body = doc.element.body

    children = list(body.iterchildren())

    if not children:

        print('Tài liệu rỗng.')

        return None

    first_element = children[0]

    if first_element.tag != qn('w:tbl'):

        print('Phần tử đầu tiên không phải bảng.')

        return None

    table = doc.tables[0]

    para_texts = []

    for row in table.rows:

        for cell in row.cells:

            for para in cell.paragraphs:

                text = para.text.strip()

                if not text:

                    continue

                para_texts.append(text)

    first_element.getparent().remove(first_element)

    for text in reversed(para_texts):

        new_p = doc.add_paragraph()

        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = new_p.add_run(text)

        run.bold = True

        run.font.color.rgb = RGBColor(0, 0, 255)

        body.insert(0, new_p._element)





def xoa_highlight_shading(doc):

    pass

# WARNING: Decompyle incomplete





def ABCD_dam_xanh(doc):

    pattern = re.compile('^([ABCD]\\. )')

    indices = tim_chi_so_phuong_an_ABCD(doc)

    for i in indices:

        para = doc.paragraphs[i]

        runs = para.runs

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(0, 0, 255)

        runs[0].font.underline = False





def abcd_nho_dam_xanh(doc):

    pattern = re.compile('^([abcd]\\) )')

    indices = tim_chi_so_phuong_an_abcd_nho(doc)

    for i in indices:

        para = doc.paragraphs[i]

        runs = para.runs

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(0, 0, 255)

        runs[0].font.underline = False





def Xoa_dong_start_with(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not paragraphs[p].text.startswith(text):

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def bo_danh_dau_dap_an(doc):

    ABCD_dam_xanh(doc)

    abcd_nho_dam_xanh(doc)

    remove_shading_highlight_from_docx(doc)

    Xoa_dong_start_with(doc, 'ĐS:')





def remove_shading_highlight_from_docx(doc):

    pass

# WARNING: Decompyle incomplete





def convert_font_to_times_new_roman(doc):

    for para in doc.paragraphs:

        for run in para.runs:

            run.font.name = 'Times New Roman'

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for run in para.runs:

                        run.font.name = 'Times New Roman'





def change_font_size(doc):

    for para in doc.paragraphs:

        for run in para.runs:

            run.font.size = Pt(12)

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for run in para.runs:

                        run.font.size = Pt(12)





def so_anh_chua_inline(doc):

    cau_first = tim_cau_first(doc)

    namespaces = {

        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',

        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing' }

    so_anh = 0

    for p in range(cau_first, len(doc.paragraphs)):

        para = doc.paragraphs[p]

        p_el = para._element

        anchors = p_el.findall('.//wp:anchor', namespaces)

        so_anh += len(anchors)

    return so_anh





def Canh_bao_anh_chua_inline(doc, messages):

    so_anh = so_anh_chua_inline(doc)

    if so_anh > 0:

        messages.append(f'''Có {so_anh} ảnh không được đặt ở chế độ inline, hãy chuyển ảnh về inline tex trước khi làm việc khác, bạn có thể mở word lên và dùng công cụ Pict inline ở Tool 1 cho nhanh.''')

        return None





def Canh_bao_da_chuyen_inline(doc, messages):

    so_anh = so_anh_chua_inline(doc)

    if so_anh > 0:

        messages.append(f'''Có {so_anh} ảnh được chuyển qua inline, hãy kiểm tra lại vị trí xuất hiện''')

        return None





def Pict_inline(doc):

    cau_first = tim_cau_first(doc)

    namespaces = {

        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',

        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing' }

    so_anh = 0

    for p in range(cau_first, len(doc.paragraphs)):

        para = doc.paragraphs[p]

        p_el = para._element

        anchors = p_el.findall('.//wp:anchor', namespaces)

        so_anh += len(anchors)

        for anchor in anchors:

            inline = etree.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')

            inline.extend(list(anchor))

            parent = anchor.getparent()

            parent.replace(anchor, inline)





def Pict_inline_thong_bao(doc, messages):

    Canh_bao_da_chuyen_inline(doc, messages)

    Pict_inline(doc)





def tim_chi_so_cau_phan(doc):

    pattern1 = '^Câu [0-9]{1,}'

    pattern2 = '^Question [0-9]{1,}'

    pattern3 = '^PHẦN I'

    indices_cau = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text

        if not re.match(pattern1, para.text, re.IGNORECASE) and re.match(pattern2, para.text, re.IGNORECASE) and re.match(pattern3, para.text, re.IGNORECASE):

            continue

        indices_cau.append(j)

    return indices_cau





def tim_chi_so_para_chua_tu_het(doc):

    """

    Trả về index của paragraph đầu tiên chứa từ 'HẾT' hoặc 'Hết'.

    Nếu không tìm thấy, trả về -1.

    """

    for i, para in enumerate(doc.paragraphs):

        text = para.text

        if not 'HẾT' in text and 'Hết' in text:

            continue

        

        return enumerate(doc.paragraphs), i

    return -1





def page_A4_1C(doc):

    '''Thiết lập trang thành 2 cột với khoảng cách phù hợp'''

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.4)

    section.bottom_margin = Inches(0.4)

    section.left_margin = Inches(0.4)

    section.right_margin = Inches(0.4)

    section.gutter = Inches(0)

    section.header_distance = Inches(0.24)

    section.footer_distance = Inches(0.24)

    end = len(doc.paragraphs)

    index_het = tim_chi_so_para_chua_tu_het(doc)

    if index_het != -1:

        end = index_het

    for p in range(end):

        para = doc.paragraphs[p]

        para.paragraph_format.first_line_indent = InchesToPoints(0)

        para.paragraph_format.left_indent = InchesToPoints(0)

        para.paragraph_format.right_indent = InchesToPoints(0)

    sectPr = doc.sections[0]._sectPr

    cols_list = sectPr.findall(qn('w:cols'))

    for c in cols_list:

        sectPr.remove(c)





def page_le_trang(doc):

    '''Thiết lập trang thành 2 cột với khoảng cách phù hợp'''

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.4)

    section.bottom_margin = Inches(0.4)

    section.left_margin = Inches(0.4)

    section.right_margin = Inches(0.4)

    section.gutter = Inches(0)

    section.header_distance = Inches(0.24)

    section.footer_distance = Inches(0.24)

    end = len(doc.paragraphs)

    index_het = tim_chi_so_para_chua_tu_het(doc)

    if index_het != -1:

        end = index_het

    for p in range(end):

        para = doc.paragraphs[p]

        para.paragraph_format.first_line_indent = InchesToPoints(0)

        para.paragraph_format.left_indent = InchesToPoints(0)

        para.paragraph_format.right_indent = InchesToPoints(0)





def xoa_pagebreak(doc):

    ns = {

        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main' }

    for para in doc.paragraphs:

        runs_to_remove = []

        for run in para.runs:

            for br in run._element.findall('.//w:br', namespaces = ns):

                if not br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':

                    continue

                runs_to_remove.append(run)

        for run in runs_to_remove:

            para._element.remove(run._element)





def xoa_section_break(doc):

    if len(doc.sections) > 1:

        for i in range(len(doc.sections) - 1):

            second_section = doc.sections[0]

            second_section._sectPr.getparent().remove(second_section._sectPr)

        return None





def delete_bookmarks(doc):

    excluded_names = {

        'MDH',

        'num_page'}

    bookmark_start_ids_to_keep = set()

    for elem in doc.element.body.iter():

        if not elem.tag == qn('w:bookmarkStart'):

            continue

        name = elem.get(qn('w:name'))

        bookmark_id = elem.get(qn('w:id'))

        if not name in excluded_names:

            continue

        bookmark_start_ids_to_keep.add(bookmark_id)

# WARNING: Decompyle incomplete





def tim_chi_so_phuong_an_ABCD(doc):

    '''para, đậm xanh, tab'''

    labels = [

        'A.',

        'B.',

        'C.',

        'D.']

    indices_phuong_an = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text.strip()

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def tim_chi_so_phuong_an_abcd_nho(doc):

    '''para, đạm xanh, tab'''

    labels = [

        'a)',

        'b)',

        'c)',

        'd)']

    indices_phuong_an = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text.strip()

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def tim_chi_so_cau(doc):

    '''para, đạm xanh, tab'''

    pattern1 = '^Question [0-9]{1,}[.:]'

    pattern2 = '^Câu [0-9]{1,}[.:]'

    indices_cau = []

    for i, para in enumerate(doc.paragraphs):

        if not re.match(pattern1, para.text) and re.match(pattern2, para.text):

            continue

        indices_cau.append(i)

    return indices_cau





def tim_chi_so_phuong_an_ABCD_abcd_xoa_tab(doc):

    """Tìm phương án trong câu hỏi theo định dạng 'A.', 'B.', 'C.', 'D.' từ from_a đến to_b"""

    labels = [

        'A.',

        'B.',

        'C.',

        'D.',

        'a)',

        'b)',

        'c)',

        'd)']

    indices_phuong_an = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text

        if not text.strip().startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    return indices_phuong_an





def xoa_tab_dau_phuong_an(doc):

    indices = tim_chi_so_phuong_an_ABCD_abcd_xoa_tab(doc)

    for p in indices:

        para = doc.paragraphs[p]

        if not para.runs:

            continue

        first_run = para.runs[0]

        if not first_run.text:

            continue

        first_run.text = re.sub('^[ \\t]+', '', first_run.text)





def replace_ABCD_space_docx(doc):

    indices_phuong_an = tim_chi_so_phuong_an_ABCD(doc)

    pattern = '[ ]{2,}[A-D]\\.'

    for p in indices_phuong_an:

        paragraph = doc.paragraphs[p]

        for i in range(4):

            full_text = (lambda .0: pass# WARNING: Decompyle incomplete

)(paragraph.runs())

            match = re.search(pattern, full_text)

            if not match:

                continue

            span_start = match.start()

            span_end = match.end()

            matched_text = match.group()

            whitespace_match = re.match('[ ]{2,}', matched_text)

            whitespace_len = len(whitespace_match.group())

            replace_start = span_start

            replace_end = span_start + whitespace_len

            current_pos = 0

            for run in paragraph.runs:

                run_len = len(run.text)

                if current_pos + run_len <= replace_start or current_pos >= replace_end:

                    current_pos += run_len

                    continue

                rel_start = max(0, replace_start - current_pos)

                rel_end = min(run_len, replace_end - current_pos)

                before = run.text[:rel_start]

                after = run.text[rel_end:]

                current_pos += run_len





def insert_paragraph_after(paragraph, doc, index):

    '''Tạo một đoạn văn mới ngay sau `paragraph` và cập nhật danh sách doc.paragraphs'''

    new_paragraph = OxmlElement('w:p')

    paragraph._element.addnext(new_paragraph)

    new_doc_para = doc.paragraphs[index]._parent.add_paragraph()

    doc.paragraphs.insert(index + 1, new_doc_para)

    return doc.paragraphs[index + 1]





def mu11_thanh_13(doc):

    i = 0

    if i < len(doc.paragraphs):

        para = doc.paragraphs[i]

        if '\n' in para.text:

            new_para = insert_paragraph_after(para, doc, i)

            move_element = False

            for child in list(para._element):

                if move_element:

                    new_para._element.append(child)

                    continue

                if not child.tag.endswith('r'):

                    continue

                if not child.text:

                    continue

                if not '\n' in child.text:

                    continue

                split_text = child.text.split('\n', 1)

                child.text = split_text[0]

                new_child = deepcopy(child)

                new_child.text = split_text[1]

                new_para._element.append(new_child)

                move_element = True

        i += 1

        if i < len(doc.paragraphs):

            continue

        return None





def split_paragraph_with_tabs(doc):

    pass

# WARNING: Decompyle incomplete





def delete_run0_empty_indices_ok(doc, indices):

    

    def is_safe_to_delete(run):

        """

        Kiểm tra xem run có an toàn để xóa không.



        Run sẽ được giữ lại nếu chứa bất kỳ thành phần nào sau:

        - Văn bản (w:t)

        - Hình ảnh (w:drawing, w:pict)

        - Công thức toán học Word (namespace 'm:')

        - Công thức MathType (w:object, v:shape, w:smartTag, OLEObject)

        """

        pass

    # WARNING: Decompyle incomplete



    for i in indices:

        para = doc.paragraphs[i]

        if not para.runs:

            continue

        run = para.runs[0]

        if not run.text.strip() == '':

            continue

        if not is_safe_to_delete(run):

            continue

        run._element.getparent().remove(run._element)





def tim_chi_so_cau_dap_an_phan_chum(doc):

    """Tìm phương án trong câu hỏi theo định dạng 'A.', 'B.', 'C.', 'D.' từ from_a đến to_b"""

    indices = []

    pattern1 = '^Câu [0-9]{1,}[.:]'

    pattern2 = '^Question [0-9]{1,}[.:]'

    labels = [

        'A.',

        'B.',

        'C.',

        'D.',

        'a)',

        'b)',

        'c)',

        'd)',

        'S1@',

        'E1@',

        'S2@',

        'E2@',

        'S3@',

        'E3@',

        'S4@',

        'E4@',

        '<S',

        '<E',

        'ĐS:']

    for p in range(len(doc.paragraphs)):

        para = doc.paragraphs[p]

        text = para.text

        if not re.match(pattern1, text) and re.match(pattern2, text) and text.startswith(tuple(labels)):

            continue

        indices.append(p)

    return indices





def delete_run0_empty(doc):

    indices = tim_chi_so_cau_dap_an_phan_chum(doc)

    delete_run0_empty_indices_ok(doc, indices)





def xoa_tab_thua_phuong_an(doc):

    

    def remove_tabs_in_paragraphs(doc, p):

        para = doc.paragraphs[p]

        i = 0

        if i < len(para.runs):

            run = para.runs[i]

            if '\t' in run.text:

                text_parts = run.text.split('\t')

                insert_index = list(para._element).index(run._element)

                prev_run = run

                run._element.getparent().remove(run._element)

                for part in reversed(text_parts):

                    if not part:

                        continue

                    new_run = para.add_run(part)

                    new_run.bold = prev_run.bold

                    new_run.italic = prev_run.italic

                    new_run.underline = prev_run.underline

                    new_run.font.color.rgb = prev_run.font.color.rgb

                    new_run.font.size = prev_run.font.size

                    para._element.insert(insert_index, new_run._element)

                    insert_index += 1

            else:

                i += 1

            if i < len(para.runs):

                continue

            return None



    indices = tim_chi_so_phuong_an_ABCD_abcd_xoa_tab(doc)

    for p in indices:

        remove_tabs_in_paragraphs(doc, p)

    indices_cau = tim_chi_so_cau(doc)

    for p in indices_cau:

        remove_tabs_in_paragraphs(doc, p)





def ABCD_dam_xanh_red(doc):

    pattern = re.compile('^([ABCD].)')

    indices = tim_chi_so_phuong_an_ABCD(doc)

    for i in indices:

        para = doc.paragraphs[i]

        runs = para.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            if not run.text:

                enumerate(runs)

            else:

                full_text += run.text

                match_run_indices.append(idx)

                if not re.match(pattern, full_text.strip()):

                    continue

                indices

    match = re.match(pattern, full_text.strip())

    if not match:

        continue

    matched_text = match.group(1)

    remaining_text = full_text[len(matched_text):]

    last_run_idx = match_run_indices[-1]

    last_run = runs[last_run_idx]

    cloned_run = deepcopy(last_run)

    if not runs[0].font.underline:

        runs[0].font.underline

        if runs[0].font.color:

            runs[0].font.color

    Dap_an = runs[0].font.color.rgb == RGBColor(255, 0, 0)

    if Dap_an:

        runs[0].text = matched_text

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(255, 0, 0)

        runs[0].font.underline = True

    else:

        runs[0].text = matched_text

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(0, 0, 255)

    for idx in reversed(match_run_indices[1:]):

        para._element.remove(runs[idx]._element)

    if not remaining_text:

        continue

    cloned_run.text = remaining_text

    runs[0]._element.addnext(cloned_run._element)

    continue





def abcd_nho_dam_xanh_red(doc):

    pattern = re.compile('^([abcd]\\))')

    indices = tim_chi_so_phuong_an_abcd_nho(doc)

    for i in indices:

        para = doc.paragraphs[i]

        runs = para.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            if not run.text:

                enumerate(runs)

            else:

                full_text += run.text

                match_run_indices.append(idx)

                if not re.match(pattern, full_text.strip()):

                    continue

                indices

    match = re.match(pattern, full_text.strip())

    if not match:

        continue

    matched_text = match.group(1)

    remaining_text = full_text[len(matched_text):]

    last_run_idx = match_run_indices[-1]

    last_run = runs[last_run_idx]

    cloned_run = deepcopy(last_run)

    if not runs[0].font.underline:

        runs[0].font.underline

        if runs[0].font.color:

            runs[0].font.color

    Dap_an = runs[0].font.color.rgb == RGBColor(255, 0, 0)

    if Dap_an:

        runs[0].text = matched_text

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(255, 0, 0)

        runs[0].font.underline = True

    else:

        runs[0].text = matched_text

        runs[0].bold = True

        runs[0].font.color.rgb = RGBColor(0, 0, 255)

    for idx in reversed(match_run_indices[1:]):

        para._element.remove(runs[idx]._element)

    if not remaining_text:

        continue

    cloned_run.text = remaining_text

    runs[0]._element.addnext(cloned_run._element)

    continue





def tim_chi_so_phan(doc):

    pattern1 = '^PHẦN I'

    pattern2 = '^PHẦN [0-9]'

    indices_phan = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text

        if not re.match(pattern1, para.text, re.IGNORECASE) and re.match(pattern2, para.text, re.IGNORECASE):

            continue

        indices_phan.append(j)

    indices_phan.append(len(doc.paragraphs))

    if len(indices_phan) > 1:

        return indices_phan

    indices_phan = [

        None] + indices_phan

    return indices_phan





def tim_chi_so_cau_from_a_to_b(doc, a, b):

    pattern1 = '\\s*(câu|question)\\s*\\d+[ \\t]*[.:]*[ \\t]*'

    indices_cau = []

    for j in range(a, b):

        para = doc.paragraphs[j]

        text = para.text

        if not re.match(pattern1, para.text, re.IGNORECASE):

            continue

        indices_cau.append(j)

    return indices_cau





def xoa_auto_number(doc):

    i = 1

# WARNING: Decompyle incomplete





def dam_xanh_cau_STT(doc):

    pattern = re.compile('\\s*(câu|question)\\s*\\d+[ \\t]*[.:]', re.IGNORECASE)

    indices_phan = tim_chi_so_phan(doc)

    if len(indices_phan) > 1:

        for k in range(len(indices_phan) - 1):

            indices_cau = tim_chi_so_cau_from_a_to_b(doc, indices_phan[k], indices_phan[k + 1])

            STT = 1

            for i in indices_cau:

                para = doc.paragraphs[i]

                runs = para.runs

                full_text = ''

                match_run_indices = []

                for idx, run in enumerate(runs):

                    if not run.text:

                        enumerate(runs)

                    else:

                        full_text += run.text

                        match_run_indices.append(idx)

                        if not re.match(pattern, full_text):

                            continue

                        indices_cau

            match = re.match(pattern, full_text)

            if not match:

                continue

            matched_text = match.group()

            tu_khoa = match.group(1).capitalize()

            remaining_text = full_text[len(matched_text):]

            last_run_idx = match_run_indices[-1]

            last_run = runs[last_run_idx]

            cloned_run = deepcopy(last_run)

            runs[0].text = f'''{tu_khoa} {STT}.'''

            runs[0].bold = True

            runs[0].font.color.rgb = RGBColor(0, 0, 255)

            for idx in reversed(match_run_indices[1:]):

                para._element.remove(runs[idx]._element)

            if remaining_text:

                cloned_run.text = remaining_text

                runs[0]._element.addnext(cloned_run._element)

            STT += 1

        continue

        return None





def dam_xanh_question_STT(doc):

    pattern = re.compile('^(Question \\d+[.:])')

    indices_phan = tim_chi_so_phan(doc)

    if len(indices_phan) > 1:

        for k in range(len(indices_phan) - 1):

            indices_cau = tim_chi_so_cau_from_a_to_b(doc, indices_phan[k], indices_phan[k + 1])

            STT = 1

            for i in indices_cau:

                para = doc.paragraphs[i]

                runs = para.runs

                full_text = ''

                match_run_indices = []

                for idx, run in enumerate(runs):

                    if not run.text:

                        enumerate(runs)

                    else:

                        full_text += run.text

                        match_run_indices.append(idx)

                        if not re.match(pattern, full_text.strip()):

                            continue

                        indices_cau

            match = re.match(pattern, full_text.strip())

            if not match:

                continue

            matched_text = match.group(1)

            remaining_text = full_text[len(matched_text):]

            last_run_idx = match_run_indices[-1]

            last_run = runs[last_run_idx]

            cloned_run = deepcopy(last_run)

            runs[0].text = f'''Question {STT}.'''

            runs[0].bold = True

            runs[0].font.color.rgb = RGBColor(0, 0, 255)

            for idx in reversed(match_run_indices[1:]):

                para._element.remove(runs[idx]._element)

            if remaining_text:

                cloned_run.text = remaining_text

                runs[0]._element.addnext(cloned_run._element)

            STT += 1

        continue

        return None





def canh_pict_giua_python_docx(doc):

    for paragraph in doc.paragraphs:

        if not 'graphic' in paragraph._element.xml:

            continue

        text = paragraph.text.strip()

        if not text == '':

            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER





def canh_deu_2ben_python_docx(doc):

    pass

# WARNING: Decompyle incomplete





def canh_before_after_all(doc):

    for p in range(len(doc.paragraphs)):

        paragraph = doc.paragraphs[p]

        paragraph.paragraph_format.space_before = Pt(0)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def tim_chi_so_dap_o(doc):

    pattern1 = '^PHẦN I'

    pattern2 = '^PHẦN [0-9]'

    indices_phan = []

    for j in range(len(doc.paragraphs)):

        para = doc.paragraphs[j]

        text = para.text

        if not re.match(pattern1, para.text, re.IGNORECASE) and re.match(pattern2, para.text, re.IGNORECASE):

            continue

        indices_phan.append(j)

    indices_phan.append(len(doc.paragraphs))

    if len(indices_phan) > 1:

        return indices_phan

    indices_phan = [

        None] + indices_phan

    return indices_phan





def chuyen_ve_DS(doc):

    keywords = [

        'đáp án:',

        'đa:',

        'đáp số:',

        'tl:',

        'trả lời:',

        'kết quả:',

        'kq:']

    for para in doc.paragraphs:

        for run in para.runs:

            text = run.text.strip()

            lower_text = text.lower()

            for keyword in keywords:

                if not lower_text.startswith(keyword):

                    continue

                run.text = run.text.strip()

                length = len(keyword)

                run.text = 'ĐS:' + run.text[length:]

                keywords





def chuanhoa_dap_so_docx(doc):

    pattern = re.compile('\\s*(đáp án|đa|đáp số|đs|trả lời|kết quả|kq)[ \\t]*[:][ \\t]*', re.IGNORECASE)

    for para in doc.paragraphs:

        runs = para.runs

        full_text = ''

        match_run_indices = []

        for idx, run in enumerate(runs):

            if not run.text:

                enumerate(runs)

            else:

                full_text += run.text

                match_run_indices.append(idx)

                if not re.match(pattern, full_text):

                    continue

                doc.paragraphs

    match = re.match(pattern, full_text)

    if not match:

        continue

    matched_text = match.group()

    tu_khoa = match.group(1).capitalize()

    remaining_text = full_text[len(matched_text):]

    last_run_idx = match_run_indices[-1]

    last_run = runs[last_run_idx]

    cloned_run = deepcopy(last_run)

    runs[0].text = 'ĐS:'

    runs[0].bold = True

    runs[0].font.name = 'Times New Roman'

    runs[0].font.size = Pt(12)

    runs[0].font.color.rgb = RGBColor(255, 0, 0)

    for idx in reversed(match_run_indices[1:]):

        para._element.remove(runs[idx]._element)

    if not remaining_text:

        continue

    cloned_run.text = remaining_text

    runs[0]._element.addnext(cloned_run._element)

    continue





def Red_Bold_DS(doc):

    for para in doc.paragraphs:

        if not para.text.strip().startswith('ĐS:'):

            continue

        if para.text.strip().endswith('.'):

            for run in reversed(para.runs):

                if not run.text.strip().endswith('.'):

                    continue

                run.text = run.text.strip().rstrip('.')

                reversed(para.runs)

        for run in para.runs:

            if not run.text:

                continue

            run.text = run.text.strip().replace(' ', '').replace('.', ',')

            run.bold = True

            run.font.name = 'Times New Roman'

            run.font.size = Pt(12)

            run.font.color.rgb = RGBColor(255, 0, 0)





def tim_chi_so_cau_dap_an_para(doc):

    indices = []

    pattern1 = '^Câu [0-9]{1,}[.:]'

    pattern2 = '^Question [0-9]{1,}[.:]'

    labels = [

        'A.',

        'B.',

        'C.',

        'D.',

        'a)',

        'b)',

        'c)',

        'd)']

    for p in range(len(doc.paragraphs)):

        para = doc.paragraphs[p]

        text = para.text

        if not re.match(pattern1, text) and re.match(pattern2, text) and text.startswith(tuple(labels)):

            continue

        indices.append(p)

    return indices





def add_space_run_0(doc):

    indices = tim_chi_so_cau_dap_an_para(doc)

    for i in indices:

        para = doc.paragraphs[i]

        if not len(para.runs) > 0:

            continue

        run = para.runs[0]

        text = run.text

        if not text:

            continue

        run.text = text.rstrip() + ' '





def xoa_space_run_1(doc):

    indices = tim_chi_so_cau_dap_an_para(doc)

# WARNING: Decompyle incomplete





def add_space_run_1(doc):

    indices = tim_chi_so_cau_dap_an_para(doc)

    for i in indices:

        para = doc.paragraphs[i]

        if not len(para.runs) > 1:

            continue

        run0 = para.runs[0]

        text0 = run0.text

        run = para.runs[1]

        text = run.text

        if text0.endswith(' '):

            continue

        if not text:

            continue

        run.text = ' ' + text.lstrip()





def xoa_space_cuoi_dong(doc):

    pass

# WARNING: Decompyle incomplete





def xuong_dong_phuong_an_chung_docx(doc):

    xoa_ky_tu_an_docx(doc)

    thay_ky_tu_gia_tao_ABCD_docx(doc)

    xoa_section_break(doc)

    xoa_pagebreak(doc)

    delete_bookmarks(doc)

    mu11_thanh_13(doc)

    page_A4_1C(doc)

    delete_run0_empty(doc)

    xoa_tab_dau_phuong_an(doc)

    for i in range(6):

        split_paragraph_with_tabs(doc)

    xoa_tab_thua_phuong_an(doc)

    delete_run0_empty(doc)

    xoa_dong_trang_docx_new(doc)

    ABCD_dam_xanh_red(doc)

    abcd_nho_dam_xanh_red(doc)

    dam_xanh_cau_STT(doc)

    canh_pict_giua_python_docx(doc)

    canh_before_after_all(doc)

    canh_deu_2ben_python_docx(doc)

    chuanhoa_dap_so_docx(doc)

    Red_Bold_DS(doc)

    add_space_run_0(doc)

    xoa_space_run_1(doc)

    xoa_space_cuoi_dong(doc)

    chuanhoa_loigiai_docx(doc)





def xuong_dong_phuong_an_docx(doc, messages):

    Canh_bao_da_chuyen_inline(doc, messages)

    Pict_inline(doc)

    if has_numbered_paragraph(doc):

        messages.append('Có một số câu được đánh số tự động, đã chuyển, hãy kiểm tra lại')

    xoa_auto_number(doc)

    xuong_dong_phuong_an_chung_docx(doc)





def xuong_dong_phuong_an_docx_tool(doc, messages):

    Canh_bao_da_chuyen_inline(doc, messages)

    Pict_inline(doc)

    if has_numbered_paragraph(doc):

        messages.append('Có một số câu được đánh số tự động, đã chuyển, hãy kiểm tra lại')

    xoa_auto_number(doc)

    xuong_dong_phuong_an_chung_docx(doc)

    replace_text_in_doc(doc, '\t', ' ')

    replace_pattern_in_doc(doc, '[ ]{2,}', ' ')





def fix_loi_sau_xd_phuong_an_docx(doc):

    xoa_ky_tu_an_docx(doc)

    thay_ky_tu_gia_tao_ABCD_docx(doc)

    xoa_section_break(doc)

    xoa_pagebreak(doc)

    delete_bookmarks(doc)

    mu11_thanh_13(doc)

    delete_run0_empty(doc)





def resize_images_to_column_width(doc):

    '''Chỉnh kích thước hình ảnh sao cho không quá 3.5 inches chiều rộng'''

    max_width = Inches(3.5)

    for shape in doc.inline_shapes:

        if not shape.width > max_width:

            continue

        aspect_ratio = shape.height / shape.width

        shape.width = int(max_width)

        shape.height = int(max_width * aspect_ratio)





def page_2_cot_Mix(doc):

    '''Thiết lập trang thành 2 cột với khoảng cách phù hợp'''

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.4)

    section.bottom_margin = Inches(0.4)

    section.left_margin = Inches(0.4)

    section.right_margin = Inches(0.4)

    section.gutter = Inches(0)

    section.header_distance = Inches(0.24)

    section.footer_distance = Inches(0.24)

    indices_cau = tim_chi_so_cau_phan(doc)

    indices_phuong_an = tim_chi_so_phuong_an_ABCD_abcd_xoa_tab(doc)

    for p in range(len(doc.paragraphs)):

        para = doc.paragraphs[p]

        para.paragraph_format.first_line_indent = InchesToPoints(0)

        para.paragraph_format.left_indent = InchesToPoints(0)

        para.paragraph_format.right_indent = InchesToPoints(0)

    sectPr = doc.sections[0]._sectPr

    cols = OxmlElement('w:cols')

    cols.set(qn('w:num'), '2')

    cols.set(qn('w:sep'), '1')

    spacing_in_twips = int(144)

    cols.set(qn('w:space'), str(spacing_in_twips))

    sectPr.append(cols)





def canh_before_after(doc, p):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_4_at(doc, p):

    tab_stops = [

        Inches(0.21),

        Inches(1.97),

        Inches(3.73),

        Inches(5.49)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_at(doc, p):

    tab_stops = [

        Inches(0.21),

        Inches(3.73)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_cot_4_at(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(0.98),

        Inches(1.86),

        Inches(2.74)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def tab_btp_2_cot_2_at(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(1.86)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def insert_tab_before_first_child(doc, p):

    tab_stops = [

        Inches(0.1),

        Inches(1.86)]

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    para = paragraphs[p]

# WARNING: Decompyle incomplete





def insert_tab_before_indices(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)





def MergeChildObjectsIntoGroup(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_1(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_1_2_cot(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_2_1_1(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_2_1_1_2_cot(doc, indices):

    pass

# WARNING: Decompyle incomplete





def option_4(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)

        tab_btp_4_at(doc, p)





def option_4_2_cot(doc, indices):

    for p in indices:

        insert_tab_before_first_child(doc, p)

        tab_btp_2_cot_4_at(doc, p)





def len_option(doc, p):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    if p < 0 or p >= len(paragraphs):

        raise IndexError('Chỉ số p không hợp lệ.')

    para = paragraphs[p]

    if not para.tag.endswith('p'):

        return 0

    text_length = 0

    equation_length = 0

    mathtype_length = 0

    image_length = 0

    nsmap = doc.element.nsmap.copy()

    if 'a' not in nsmap:

        nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    if 'm' not in nsmap:

        nsmap['m'] = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    if 'o' not in nsmap:

        nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

    for run in para.findall('.//w:r', namespaces = nsmap):

        for t in run.findall('.//w:t', namespaces = nsmap):

            if not t.text:

                continue

            text_length += len(t.text)

    math_objects = para.findall('.//m:oMathPara', namespaces = nsmap)

    if not math_objects:

        math_objects = para.findall('.//m:oMath', namespaces = nsmap)

    for math in math_objects:

        for mt in math.findall('.//m:t', namespaces = nsmap):

            if not mt.text:

                continue

            equation_length += len(mt.text)

    equation_length = int(1.7 * equation_length)

# WARNING: Decompyle incomplete





def max_len_option(doc, indices):

    max_length = 0

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for p in indices:

        if p < 0 or p >= len(paragraphs):

            raise IndexError('Chỉ số p không hợp lệ.')

        current_length = len_option(doc, p)

        if not current_length > max_length:

            continue

        max_length = current_length

    return max_length





def tim_cau_trong_all(doc):

    """Tìm tất cả vị trí chứa 'Question X.' hoặc 'Câu X.' trong tài liệu"""

    pattern1 = '^Question [0-9]{1,}[.:]'

    pattern2 = '^Câu [0-9]{1,}[.:]'

    indices_cau = []

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i, para in enumerate(paragraphs):

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        if not re.match(pattern1, paragraph.text) and re.match(pattern2, paragraph.text):

            continue

        indices_cau.append(i)

    indices_cau.append(len(body) - 1)

    return indices_cau





def tim_phuong_an_trong_cau_phan_I(doc, from_a, to_b):

    """Tìm phương án trong câu hỏi theo định dạng 'A.', 'B.', 'C.', 'D.' từ from_a đến to_b"""

    labels = [

        'A.',

        'B.',

        'C.',

        'D.']

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    indices_phuong_an = []

    for j in range(from_a, min(to_b, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    indices_phuong_an.append(to_b)

    return indices_phuong_an





def tim_phuong_an_trong_cau_phan_II(doc, from_a, to_b):

    """Tìm phương án trong câu hỏi theo định dạng 'a)', 'b)', 'c)', 'd)' từ from_a đến to_b"""

    labels = [

        'a)',

        'b)',

        'c)',

        'd)']

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    indices_phuong_an = []

    for j in range(from_a, min(to_b, len(paragraphs))):

        para = paragraphs[j]

        if not para.tag.endswith('p'):

            continue

        paragraph = Paragraph(para, doc)

        text = paragraph.text

        if not text.startswith(tuple(labels)):

            continue

        indices_phuong_an.append(j)

    indices_phuong_an.append(to_b)

    return indices_phuong_an





def lay_size_font(doc, index):

    '''Lấy kích thước font của văn bản đầu tiên trong đoạn văn tại index.'''

    if index < len(doc.paragraphs):

        paragraph = doc.paragraphs[index]

        for run in paragraph.runs:

            if not run.font:

                continue

            if not run.font.size:

                continue

            

            return paragraph.runs, run.font.size.pt

    return 12





def sap_xep_phan_I_trong_cac_vung(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 25:

            option_1(doc, phuong_an)

            continue

        if L1max < 49:

            option_2_1_1(doc, phuong_an)

            continue

        option_4(doc, phuong_an)





def sap_xep_phan_I_trong_cac_vung_font13(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 23:

            option_1(doc, phuong_an)

            continue

        if L1max < 45:

            option_2_1_1(doc, phuong_an)

            continue

        option_4(doc, phuong_an)





def sap_xep_phan_I_trong_cac_vung_font14(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 21:

            option_1(doc, phuong_an)

            continue

        if L1max < 41:

            option_2_1_1(doc, phuong_an)

            continue

        option_4(doc, phuong_an)





def sap_xep_phan_I(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_trong_cac_vung_font13(doc)

        return None

    if size_font == 14:

        sap_xep_phan_I_trong_cac_vung_font14(doc)

        return None

    sap_xep_phan_I_trong_cac_vung(doc)





def sap_xep_phan_I_2_cot_trong_cac_vung(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 12:

            option_1_2_cot(doc, phuong_an)

            continue

        if L1max < 23:

            option_2_1_1_2_cot(doc, phuong_an)

            continue

        option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot_trong_cac_vung_font13(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 11:

            option_1_2_cot(doc, phuong_an)

            continue

        if L1max < 21:

            option_2_1_1_2_cot(doc, phuong_an)

            continue

        option_4_2_cot(doc, phuong_an)





def sap_xep_phan_I_2_cot_trong_cac_vung_font14(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_I(doc, cau[k], cau[k + 1])

        if not len(phuong_an) > 2:

            continue

        phuong_an = phuong_an[:-1]

        L1max = max_len_option(doc, phuong_an)

        if L1max < 10:

            option_1_2_cot(doc, phuong_an)

            continue

        if L1max < 19:

            option_2_1_1_2_cot(doc, phuong_an)

            continue

        option_4_2_cot(doc, phuong_an)





def tim_cau_first(doc):

    """Tìm chỉ mục đoạn văn đầu tiên chứa 'Câu X.' hoặc 'Câu X:'"""

    pattern1 = '^Câu [0-9]{1,}[.:]'

    for index, paragraph in enumerate(doc.paragraphs):

        if not re.match(pattern1, paragraph.text):

            continue

        

        return enumerate(doc.paragraphs), index

    return 0





def lay_size_font(doc, index):

    '''Lấy kích thước font của văn bản đầu tiên trong đoạn văn tại index.'''

    if index < len(doc.paragraphs):

        paragraph = doc.paragraphs[index]

        for run in paragraph.runs:

            if not run.font:

                continue

            if not run.font.size:

                continue

            

            return paragraph.runs, run.font.size.pt

    return 12





def sap_xep_phan_I_2_cot(doc):

    index = tim_cau_first(doc)

    size_font = lay_size_font(doc, index)

    if size_font == 13:

        sap_xep_phan_I_2_cot_trong_cac_vung_font13(doc)

        return None

    if size_font == 14:

        sap_xep_phan_I_2_cot_trong_cac_vung_font14(doc)

        return None

    sap_xep_phan_I_2_cot_trong_cac_vung(doc)





def lam_dep_phan_II(doc):

    cau = tim_cau_trong_all(doc)

    for k in range(len(cau) - 2, -1, -1):

        phuong_an = tim_phuong_an_trong_cau_phan_II(doc, cau[k], cau[k + 1])

        phuong_an = phuong_an[:-1]

        for p in phuong_an:

            insert_tab_before_first_child(doc, p)

            tab_btp_4_at(doc, p)





def Xoa_dong_chua_tu(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not text in paragraphs[p].text:

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def Xoa_dong_start_with(doc, text):

    paragraphs = list(doc.paragraphs)

    for p in range(len(paragraphs) - 1, -1, -1):

        if not paragraphs[p].text.startswith(text):

            continue

        p_element = paragraphs[p]._element

        p_element.getparent().remove(p_element)





def replace_text_in_doc(doc, old_text, new_text):

    '''Thay thế old_text bằng new_text trong toàn bộ tài liệu, giữ nguyên định dạng'''

    replace_in_paragraphs(doc.paragraphs, old_text, new_text)

    replace_in_tables(doc.tables, old_text, new_text)

    for section in doc.sections:

        replace_in_paragraphs(section.header.paragraphs, old_text, new_text)

        replace_in_tables(section.header.tables, old_text, new_text)

        replace_in_paragraphs(section.footer.paragraphs, old_text, new_text)

        replace_in_tables(section.footer.tables, old_text, new_text)





def replace_in_paragraphs(paragraphs, old_text, new_text):

    '''Thay thế văn bản trong danh sách đoạn văn'''

    for para in paragraphs:

        for run in para.runs:

            if not old_text in run.text:

                continue

            run.text = run.text.replace(old_text, new_text)





def replace_in_tables(tables, old_text, new_text):

    '''Thay thế văn bản trong tất cả các bảng'''

    for table in tables:

        for row in table.rows:

            for cell in row.cells:

                replace_in_paragraphs(cell.paragraphs, old_text, new_text)

                replace_in_tables(cell.tables, old_text, new_text)





def replace_pattern_in_doc(doc, pattern, new_text):

    '''Thay thế old_text bằng new_text trong toàn bộ tài liệu, giữ nguyên định dạng'''

    replace_pattern_in_paragraphs(doc.paragraphs, pattern, new_text)

    replace_pattern_in_tables(doc.tables, pattern, new_text)

    for section in doc.sections:

        replace_pattern_in_paragraphs(section.header.paragraphs, pattern, new_text)

        replace_pattern_in_tables(section.header.tables, pattern, new_text)

        replace_pattern_in_paragraphs(section.footer.paragraphs, pattern, new_text)

        replace_pattern_in_tables(section.footer.tables, pattern, new_text)





def replace_pattern_in_paragraphs(paragraphs, pattern, new_text):

    '''Thay thế văn bản trong danh sách đoạn văn'''

    for para in paragraphs:

        for run in para.runs:

            if not re.search(pattern, run.text):

                continue

            run.text = re.sub(pattern, new_text, run.text)





def replace_pattern_in_tables(tables, old_text, new_text):

    '''Thay thế văn bản trong tất cả các bảng'''

    for table in tables:

        for row in table.rows:

            for cell in row.cells:

                replace_pattern_in_paragraphs(cell.paragraphs, old_text, new_text)

                replace_pattern_in_tables(cell.tables, old_text, new_text)





def canh_before_after_for_cau(doc):

    for p in range(len(doc.paragraphs)):

        paragraph = doc.paragraphs[p]

        paragraph.paragraph_format.space_before = Pt(0)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def them_cau_cuoi(doc):

    paragraph = doc.add_paragraph()

    paragraph.add_run('Câu 00. @')





def xoa_gach_chan_tab_trong_docx_OLD(doc):

    pass

# WARNING: Decompyle incomplete





def xoa_gach_chan_tab_trong_docx(doc):

    pattern = re.compile('^\\t')

    for para in doc.paragraphs:

        runs = list(para.runs)

        for run in runs:

            text = run.text

            if not pattern.match(text):

                continue

            insert_index = list(para._element).index(run._element)

            remaining_text = text[1:]

            new_run = para.add_run('\t')

            para._element.insert(insert_index, new_run._element)





def save_and_close_docx(doc, output_file):

    doc.save(output_file)





def chuanhoaA4(doc):

    sap_xep_phan_I(doc)

    lam_dep_phan_II(doc)

    Xoa_dong_start_with(doc, 'Câu 00. @')

    xoa_dong_trang_docx_new(doc)

    return None

# WARNING: Decompyle incomplete





def chuanhoadocx(doc, messages):

    xuong_dong_phuong_an_docx_tool(doc, messages)

    chuanhoaA4(doc)

    xoa_gach_chan_tab_trong_docx(doc)

    dinh_dang_cac_para_dau_cham(doc)





def chuanhoa_bo_DA_docx(doc, messages):

    xuong_dong_phuong_an_docx_tool(doc, messages)

    bo_danh_dau_dap_an(doc)

    xoa_loi_giai_lay_de(doc, selected_headings = [])

    chuanhoaA4(doc)





def chuanhoa_2C(doc):

    sap_xep_phan_I_2_cot(doc)

    page_2_cot_Mix(doc)

    resize_images_to_column_width(doc)

    Xoa_dong_start_with(doc, 'Câu 00. @')

    xoa_dong_trang_docx_new(doc)

    return None

# WARNING: Decompyle incomplete





def chuanhoadocx_2_cot(doc, messages):

    xuong_dong_phuong_an_docx_tool(doc, messages)

    chuanhoa_2C(doc)

    xoa_gach_chan_tab_trong_docx(doc)

    dinh_dang_cac_para_dau_cham(doc)





def chuanhoa_bo_DA_2_cot(doc, messages):

    xuong_dong_phuong_an_docx_tool(doc, messages)

    bo_danh_dau_dap_an(doc)

    xoa_loi_giai_lay_de(doc, selected_headings = [])

    chuanhoa_2C(doc)





def tao_ten_moi(input_doc, name_add):

    '''Đổi tên file từ abc.docx thành abc_new.docx'''

    (dir_name, file_name) = os.path.split(input_doc)

    (name, ext) = os.path.splitext(file_name)

    new_file = f'''{name}_{name_add}{ext}'''

    return os.path.join(dir_name, new_file)





def save_file_input(doc, input_file):

    doc.save(input_file)





def save_file_HS(doc, input_file):

    output_doc = tao_ten_moi(input_file, 'HS')

    doc.save(output_doc)





def save_file_GV(doc, input_file):

    output_doc = tao_ten_moi(input_file, 'GV')

    doc.save(output_doc)





def save_file_goc_Mix(doc, input_file):

    output_doc = tao_ten_moi(input_file, 'file_goc')

    doc.save(output_doc)





def save_file_new(doc, input_file):

    new = simpledialog.askstring('Nhập tên', 'Tên file = file cũ + :')

    output_doc = tao_ten_moi(input_file, new)

    doc.save(output_doc)





def is_two_columns(doc):

    section = doc.sections[0]

    sectPr = section._sectPr

    for child in sectPr.iter():

        if not child.tag.endswith('cols'):

            continue

        num = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num')

        if not num:

            continue

        if not int(num) >= 2:

            continue

        sectPr.iter()

        return True

    return False





def sum_line(doc, from_idx, to_idx):

    sum_len = 0

    so_dong = 0

    for p in range(from_idx, to_idx):

        so_dong += len_option(doc, p) // 100 + 1

    return so_dong





def thay_the_loi_giai_bang_dong_cham_vung(doc, from_idx, to_idx, page_kind = ('A4',)):

    body = doc.element.body

    elements = list(body.iterchildren())

    from_idx = max(0, from_idx)

    to_idx = min(to_idx, len(elements) - 1)

    so_dong = sum_line(doc, from_idx, to_idx)

    for i in range(to_idx - 1, from_idx - 1, -1):

        body.remove(elements[i])

    for _ in range(so_dong):

        p = OxmlElement('w:p')

        r = OxmlElement('w:r')

        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')

        rFonts.set(qn('w:ascii'), 'Times New Roman')

        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

        sz = OxmlElement('w:sz')

        sz.set(qn('w:val'), '16')

        szCs = OxmlElement('w:szCs')

        szCs.set(qn('w:val'), '16')

        color = OxmlElement('w:color')

        color.set(qn('w:val'), '0000FF')

        rPr.extend([

            rFonts,

            sz,

            szCs,

            color])

        t = OxmlElement('w:t')

        r.append(rPr)

        r.append(t)

        p.append(r)

        paragraph = Paragraph(p, doc)

        paragraph.paragraph_format.space_before = Pt(8)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def thay_loi_giai_bang_dong_cham_de(doc, page_kind, selected_headings = ('A4', [])):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    a = tim_cau_body_all(doc, selected_headings)

    b = tim_loi_giai_body_all(doc)

    c = tao_danh_sach_loigiai_cau(a, b)

    if len(b) == len(c):

        if page_kind == '2C':

            page_2_cot_Mix(doc)

        else:

            page_A4_1C(doc)

        for start, end in reversed(list(zip(b, c))):

            thay_the_loi_giai_bang_dong_cham_vung(doc, start, end, page_kind)

    xoa_dong_trang_docx_new(doc)





def insert_paragraph_at(doc, index, text):

    body = doc._body._element

    paras = doc.paragraphs

    temp_para = doc.add_paragraph()

    run = temp_para.add_run(text)

    font = run.font

    font.name = 'Times New Roman'

    font.size = Pt(12)

    font.color.rgb = RGBColor(0, 0, 255)

    insert_point = list(body).index(paras[index]._element) if index < len(paras) else len(body)

    body.insert(insert_point, temp_para._element)





def chen_dong_cham_at(doc, from_idx, so_dong, page_kind = ('A4',)):

    body = doc.element.body

    elements = list(body.iterchildren())

    for _ in range(so_dong):

        p = OxmlElement('w:p')

        r = OxmlElement('w:r')

        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')

        rFonts.set(qn('w:ascii'), 'Times New Roman')

        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

        sz = OxmlElement('w:sz')

        sz.set(qn('w:val'), '24')

        szCs = OxmlElement('w:szCs')

        szCs.set(qn('w:val'), '24')

        color = OxmlElement('w:color')

        color.set(qn('w:val'), '0000FF')

        rPr.extend([

            rFonts,

            sz,

            szCs,

            color])

        t = OxmlElement('w:t')

        r.append(rPr)

        r.append(t)

        p.append(r)

        body.insert(from_idx, p)

        from_idx += 1

        paragraph = Paragraph(p, doc)

        paragraph.paragraph_format.space_before = Pt(8)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def dinh_dang_cac_para_dau_cham(doc):

    for para in doc.paragraphs:

        text = para.text.strip()

        if not re.match('^…{5,}', text):

            continue

        para.paragraph_format.space_before = Pt(8)

        para.paragraph_format.space_after = Pt(0)

        para.paragraph_format.line_spacing = 1.15

        for run in para.runs:

            run.font.name = 'Times New Roman'

            run.font.size = Pt(8)





def them_dong_cham_de(doc, so_dong, page_kind, selected_headings = (5, 'A4', None)):

    page_le_trang(doc)

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    a = tim_cau_body_all(doc, selected_headings)

    if len(a) > 1:

        for idx in reversed(a[1:]):

            chen_dong_cham_at(doc, idx, so_dong, page_kind)

        dinh_dang_cac_para_dau_cham(doc)

        xoa_dong_trang_docx_new(doc)

        return None





def browse_file_entry(entry_file_path):

    input_file = filedialog.askopenfilename(filetypes = [

        ('Word files', '*.docx')])

    if not input_file:

        return None

    input_file = os.path.normpath(input_file)

    entry_file_path.delete(0, tk.END)

    entry_file_path.insert(tk.END, input_file)





def browse_file(listbox_tach_de):

    files = filedialog.askopenfilenames(filetypes = [

        ('Word Files', '*.docx')])

    if not files:

        return None

    for file in files:

        input_file = os.path.normpath(file)

        listbox_tach_de.insert(tk.END, input_file)





def clear_files(listbox_tach_de):

    listbox_tach_de.delete(0, tk.END)





def Tach_de_chuan_hoa(root):

    pass

# WARNING: Decompyle incomplete





def fix_loi_dap_an_sau_xd_docx_mix():

    pass

# WARNING: Decompyle incomplete





def xuongdong_docx_mix():

    pass

# WARNING: Decompyle incomplete





def xoa_body_from_a_to_b(doc, a, b):

    body = doc.element.body

    paragraphs = list(body.iterchildren())

    for i in range(b - 1, a - 1, -1):

        body.remove(paragraphs[i])

    return None

# WARNING: Decompyle incomplete





def phan_chia_n_phan(input_file, n):

    original_doc = Document(input_file)

    body_elements = list(original_doc.element.body)

    total_len = len(body_elements) - 1

    base_dir = os.path.dirname(input_file)

    base_name = os.path.splitext(os.path.basename(input_file))[0]

    part_sizes = [

        total_len // n] * n

    for i in range(total_len % n):

        pass

    [] = None

    start = 0

    for size in part_sizes:

        end = start + size

        indices.append([

            start,

            end])

        start = end

    for i, ds in enumerate(indices):

        a = ds[0]

        b = ds[1]

        doc = Document(input_file)

        xoa_body_from_a_to_b(doc, b, total_len)

        xoa_body_from_a_to_b(doc, 0, a)

        save_path = os.path.join(base_dir, f'''{base_name}_phan_{i + 1}.docx''')

        doc.save(save_path)





def Tach_file_large(root):

    pass

# WARNING: Decompyle incomplete





def Tim_kiem_cau_hoi(root):

    pass

# WARNING: Decompyle incomplete





def Thay_the(root):

    pass

# WARNING: Decompyle incomplete



