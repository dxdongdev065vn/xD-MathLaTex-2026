import os

import tkinter as tk

from tkinter import ttk, Label, filedialog, Entry, Button, Frame, Listbox, Scrollbar, messagebox

from docx import Document

from docx.text.paragraph import Paragraph

from docx.shared import RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Inches, Pt

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from lxml import etree



ElementTree

from xml.etree.ElementTree import QName

import xml.etree.ElementTree, etree

import re

import zipfile

import bisect

from mixing import mix_docx_func as mdocx_f

from mixing import mix_docx_en as mdocx_EN

from formatting import chuan_hoa_docx as chdocx

from tools import tool_by_docx as docxtool



def InchesToPoints(inches):

    return inches * 72





def LinesToPoints(lines):

    return lines * 12





def open_doc_off_python_docx(doc_path):

    doc = Document(doc_pat)

    return doc





def remove_last_element(lst):

    return lst[:-1]





def resize_images_to_column_width(doc):

    '''Chỉnh kích thước hình ảnh sao cho không quá 3.5 inches chiều rộng'''

    max_width = Inches(3.5)

    for shape in doc.inline_shapes:

        if not shape.width > max_width:

            continue

        aspect_ratio = shape.height / shape.width

        shape.width = int(max_width)

        shape.height = int(max_width * aspect_ratio)





def page_2_cot_Mix(doc):

    '''Thiết lập trang thành 2 cột với khoảng cách phù hợp'''

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.4)

    section.bottom_margin = Inches(0.4)

    section.left_margin = Inches(0.4)

    section.right_margin = Inches(0.4)

    section.gutter = Inches(0)

    section.header_distance = Inches(0.24)

    section.footer_distance = Inches(0.24)

    sectPr = doc.sections[0]._sectPr

    cols = OxmlElement('w:cols')

    cols.set(qn('w:num'), '2')

    cols.set(qn('w:sep'), '1')

    spacing_in_twips = int(144)

    cols.set(qn('w:space'), str(spacing_in_twips))

    sectPr.append(cols)





def canh_before_after_all(doc):

    for p in range(len(doc.paragraphs)):

        paragraph = doc.paragraphs[p]

        paragraph.paragraph_format.space_before = Pt(0)

        paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.line_spacing = 1.15





def canh_before_phan(doc):

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip().lower()

        if not text.startswith('phần i'):

            continue

        paragraph.paragraph_format.space_before = Pt(6)





def canh_deu_2ben_python_docx(doc):

    pass

# WARNING: Decompyle incomplete





def tab_4_for_para(paragraph):

    '''Đặt 4 tab stop cho paragraph được truyền vào'''

    tab_stops = [

        Inches(0.21),

        Inches(1.97),

        Inches(3.73),

        Inches(5.49)]

    paragraph.paragraph_format.tab_stops.clear()

    for pos in tab_stops:

        paragraph.paragraph_format.tab_stops.add_tab_stop(pos)

    return None

# WARNING: Decompyle incomplete





def tab_2_for_para(paragraph):

    '''Đặt 4 tab stop cho paragraph được truyền vào'''

    tab_stops = [

        Inches(0.21),

        Inches(3.73)]

    paragraph.paragraph_format.tab_stops.clear()

    for pos in tab_stops:

        paragraph.paragraph_format.tab_stops.add_tab_stop(pos)

    return None

# WARNING: Decompyle incomplete





def chinh_le_trang(doc):

    sections = doc.sections

    for section in sections:

        section.page_width = Inches(8.27)

        section.page_height = Inches(11.69)

        section.top_margin = Inches(0.5)

        section.bottom_margin = Inches(0.5)

        section.left_margin = Inches(0.5)

        section.right_margin = Inches(0.5)

        section.gutter = Inches(0)

        section.header_distance = Inches(0.24)

        section.footer_distance = Inches(0.24)





def xoa_vien_bang(table):

    for row in table.rows:

        for cell in row.cells:

            tc = cell._tc

            tcPr = tc.get_or_add_tcPr()

            borders = OxmlElement('w:tcBorders')

            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):

                border = OxmlElement(f'''w:{border_name}''')

                border.set(qn('w:val'), 'nil')

                borders.append(border)

            tcPr.append(borders)





def set_cell_border(cell, **kwargs):

    '''

    Set cell\'s border

    Usage:

        set_cell_border(cell,

            top={"sz": 12, "val": "single", "color": "000000"},

            bottom={"sz": 12, "val": "single", "color": "000000"},

            start={"sz": 12, "val": "single", "color": "000000"},

            end={"sz": 12, "val": "single", "color": "000000"},

        )

    '''

    tc = cell._tc

    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn('w:tcBorders'))

# WARNING: Decompyle incomplete





def convert_font_to_times_new_roman(doc):

    for para in doc.paragraphs:

        for run in para.runs:

            run.font.name = 'Times New Roman'

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for run in para.runs:

                        run.font.name = 'Times New Roman'





def change_font_size(doc):

    for para in doc.paragraphs:

        for run in para.runs:

            run.font.size = Pt(12)

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for run in para.runs:

                        run.font.size = Pt(12)





def tao_2_bang_header(doc):

    table1 = doc.add_table(rows = 1, cols = 2)

    table1.autofit = False

    table1.cell(0, 0).width = Inches(3.75)

    table1.cell(0, 1).width = Inches(3.75)

    p1 = table1.cell(0, 0).paragraphs[0]

    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1.paragraph_format.space_before = Pt(0)

    p1.paragraph_format.space_after = Pt(0)

    p1.paragraph_format.line_spacing = 1.15

    run = p1.add_run('TRƯỜNG THPT CHUYÊN QUỐC HỌC – HUẾ')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    p1 = table1.cell(0, 0).add_paragraph()

    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1.paragraph_format.space_before = Pt(0)

    p1.paragraph_format.space_after = Pt(0)

    p1.paragraph_format.line_spacing = 1.15

    run = p1.add_run('TỔ TOÁN')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    p1 = table1.cell(0, 0).add_paragraph()

    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1.paragraph_format.space_before = Pt(0)

    p1.paragraph_format.space_after = Pt(0)

    p1.paragraph_format.line_spacing = 1.15

    run = p1.add_run('ĐỀ CHÍNH THỨC')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    p1 = table1.cell(0, 0).add_paragraph()

    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1.paragraph_format.space_before = Pt(0)

    p1.paragraph_format.space_after = Pt(0)

    p1.paragraph_format.line_spacing = 1.15

    run = p1.add_run('(Đề thi có <sotrang> trang)')

    run.italic = True

    p2 = table1.cell(0, 1).paragraphs[0]

    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2.paragraph_format.space_before = Pt(0)

    p2.paragraph_format.space_after = Pt(0)

    p2.paragraph_format.line_spacing = 1.15

    run = p2.add_run('ĐỀ KIỂM TRA GIỮA KỲ II')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    p2 = table1.cell(0, 1).add_paragraph()

    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2.paragraph_format.space_before = Pt(0)

    p2.paragraph_format.space_after = Pt(0)

    p2.paragraph_format.line_spacing = 1.15

    run = p2.add_run('NĂM HỌC 2025 – 2026')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    p2 = table1.cell(0, 1).add_paragraph()

    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2.paragraph_format.space_before = Pt(0)

    p2.paragraph_format.space_after = Pt(0)

    p2.paragraph_format.line_spacing = 1.15

    run = p2.add_run('Môn: Toán – Lớp 10')

    run.italic = True

    p2 = table1.cell(0, 1).add_paragraph()

    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2.paragraph_format.space_before = Pt(0)

    p2.paragraph_format.space_after = Pt(0)

    p2.paragraph_format.line_spacing = 1.15

    run = p2.add_run('Thời gian: 90 phút (Không kể thời gian phát đề)')

    run.italic = True

    dong_trang = doc.add_paragraph('')

    table2 = doc.add_table(rows = 1, cols = 2)

    table2.autofit = False

    table2.cell(0, 0).width = Inches(5.7)

    table2.cell(0, 1).width = Inches(1.8)

    p3 = table2.cell(0, 0).paragraphs[0]

    p3.paragraph_format.space_before = Pt(0)

    p3.paragraph_format.space_after = Pt(0)

    p3.paragraph_format.line_spacing = 1.15

    run = p3.add_run('Họ và tên học sinh:........................................................')

    run = p3.add_run('Số báo danh:.......................')

    p4 = table2.cell(0, 1).paragraphs[0]

    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4.paragraph_format.space_before = Pt(0)

    p4.paragraph_format.space_after = Pt(0)

    p4.paragraph_format.line_spacing = 1.15

    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p4.add_run('Mã đề thi: <made>')

    run.bold = True

    set_cell_border(table2.cell(0, 1), top = {

        'sz': 12,

        'val': 'single',

        'color': '000000' }, bottom = {

        'sz': 12,

        'val': 'single',

        'color': '000000' }, start = {

        'sz': 12,

        'val': 'single',

        'color': '000000' }, end = {

        'sz': 12,

        'val': 'single',

        'color': '000000' })





def tao_header_2c(doc):

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('TRƯỜNG THPT CHUYÊN QUỐC HỌC – HUẾ')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('TỔ TOÁN')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('ĐỀ KIỂM TRA GIỮA KỲ II')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('NĂM HỌC 2025 – 2026')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('Môn: Toán – Lớp 10')

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('(Thời gian làm bài 90 phút)')

    run.italic = True

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('Mã đề thi: <made>')

    run.bold = True

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run('Đề thi có <sotrang> trang')

    run.italic = True





def tao_phan_than(doc):

    para = doc.add_paragraph('')

    run = para.add_run('PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    run = para.add_run(' Thí sinh trả lời từ câu 1 đến câu 12. Mỗi câu hỏi thí sinh chỉ chọn một phương án.')

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    run = para.add_run('S1@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('Câu 1.')

    run.bold = True

    para.add_run(' Chọn đáp án đúng')

    para = doc.add_paragraph('')

    tab_4_for_para(para)

    para.add_run('\t')

    run = para.add_run('A.')

    run.bold = True

    para.add_run(' Đáp án 1.')

    para.add_run('\t')

    run = para.add_run('B.')

    run.bold = True

    para.add_run(' Đáp án 2.')

    para.add_run('\t')

    run = para.add_run('C')

    run.bold = True

    run.underline = True

    run = para.add_run('.')

    run.bold = True

    para.add_run(' Đáp án 3.')

    para.add_run('\t')

    run = para.add_run('D.')

    run.bold = True

    para.add_run(' Đáp án 4.')

    para = doc.add_paragraph('')

    run = para.add_run('Câu 2.')

    run.bold = True

    para.add_run(' Chọn đáp án đúng')

    para = doc.add_paragraph('')

    tab_2_for_para(para)

    run = para.add_run('A.')

    run.bold = True

    para.add_run(' Đáp án 1.')

    para.add_run('\t')

    run = para.add_run('B.')

    run.bold = True

    para.add_run(' Đáp án 2.')

    para = doc.add_paragraph('')

    tab_2_for_para(para)

    run = para.add_run('C')

    run.bold = True

    run.underline = True

    run = para.add_run('.')

    run.bold = True

    para.add_run(' Đáp án 3.')

    para.add_run('\t')

    run = para.add_run('D.')

    run.bold = True

    para.add_run(' Đáp án 4.')

    para = doc.add_paragraph('')

    run = para.add_run('Câu 3.')

    run.bold = True

    para.add_run(' Chọn đáp án đúng')

    para = doc.add_paragraph('')

    run = para.add_run('A.')

    run.bold = True

    para.add_run(' Đáp án 1.')

    para = doc.add_paragraph('')

    run = para.add_run('B.')

    run.bold = True

    para.add_run(' Đáp án 2.')

    para = doc.add_paragraph('')

    run = para.add_run('C')

    run.bold = True

    run.underline = True

    run = para.add_run('.')

    run.bold = True

    para.add_run(' Đáp án 3.')

    para = doc.add_paragraph('')

    run = para.add_run('D.')

    run.bold = True

    para.add_run(' Đáp án 4.')

    para = doc.add_paragraph('')

    run = para.add_run('E1@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('PHẦN II. Câu trắc nghiệm đúng sai.')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    run = para.add_run(' Thí sinh trả lời từ câu 1 đến câu 3. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.')

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    run = para.add_run('S2@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('Câu 1.')

    run.bold = True

    para.add_run(' Nội dung câu:')

    para = doc.add_paragraph('')

    run = para.add_run('a')

    run.bold = True

    run.underline = True

    run = para.add_run(')')

    run.bold = True

    para.add_run(' Phương án 1.')

    para = doc.add_paragraph('')

    run = para.add_run('b)')

    run.bold = True

    para.add_run(' Phương án 2.')

    para = doc.add_paragraph('')

    run = para.add_run('c')

    run.bold = True

    run.underline = True

    run = para.add_run(')')

    run.bold = True

    para.add_run(' Phương án 3.')

    para = doc.add_paragraph('')

    run = para.add_run('d)')

    run.bold = True

    para.add_run(' Phương án 4.')

    para = doc.add_paragraph('')

    run = para.add_run('E2@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('PHẦN III. Câu trắc nghiệm trả lời ngắn.')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    run = para.add_run(' Thí sinh trả lời từ câu 1 đến câu 4.')

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    run = para.add_run('S3@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('Câu 1.')

    run.bold = True

    para.add_run(' Nội dung câu hỏi.')

    doc.add_paragraph('ĐS:5')

    para = doc.add_paragraph('')

    run = para.add_run('E3@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('PHẦN IV. Câu hỏi tự luận.')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 0, 255)

    para = doc.add_paragraph('')

    run = para.add_run('S4@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    run = para.add_run('Câu 1.')

    run.bold = True

    para.add_run(' Nội dung câu hỏi.')

    para = doc.add_paragraph('')

    run = para.add_run('a)')

    run.bold = True

    para.add_run(' Nội dung ý 1.')

    para = doc.add_paragraph('')

    run = para.add_run('b)')

    run.bold = True

    para.add_run(' Nội dung ý 2.')

    para = doc.add_paragraph('')

    run = para.add_run('E4@')

    run.bold = True

    run.font.color.rgb = RGBColor(0, 128, 0)

    para = doc.add_paragraph('')

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para.add_run('------------ ')

    run = para.add_run('HẾT')

    run.bold = True

    para.add_run(' -------------')

    para = doc.add_paragraph('- Thí sinh không được sử dụng tài liệu.')

    para.paragraph_format.first_line_indent = Inches(1)

    para.paragraph_format.left_indent = Inches(0)

    para.paragraph_format.right_indent = Inches(0)

    para = doc.add_paragraph('- Cán bộ coi thi không giải thích gì thêm.')

    para.paragraph_format.first_line_indent = Inches(1)

    para.paragraph_format.left_indent = Inches(0)

    para.paragraph_format.right_indent = Inches(0)





def insert_field_code_ok(paragraph, field_code):

    '''Chèn field (như PAGE hoặc NUMPAGES) vào một đoạn văn bản'''

    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')

    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')

    instrText.text = field_code

    fldChar2 = OxmlElement('w:fldChar')

    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')

    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)

    run._r.append(instrText)

    run._r.append(fldChar2)

    run._r.append(fldChar3)





def insert_field_code(paragraph, field_code):

    '''Chèn field vào paragraph với định dạng font'''

    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')

    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')

    instrText.set(qn('xml:space'), 'preserve')

    instrText.text = field_code

    fldChar2 = OxmlElement('w:fldChar')

    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')

    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)

    run._r.append(instrText)

    run._r.append(fldChar2)

    run._r.append(fldChar3)

    font = run.font

    font.name = 'Times New Roman'

    font.size = Pt(11)

    font.color.rgb = RGBColor(0, 0, 0)

    font.italic = True





def set_paragraph_top_border(paragraph, color, size, space = ('auto', '4', '1')):

    '''Thêm viền trên (top border) cho paragraph'''

    p = paragraph._p

    pPr = p.get_or_add_pPr()

    pbdr = OxmlElement('w:pBdr')

    top = OxmlElement('w:top')

    top.set(qn('w:val'), 'single')

    top.set(qn('w:sz'), size)

    top.set(qn('w:space'), space)

    top.set(qn('w:color'), color)

    pbdr.append(top)

    pPr.append(pbdr)





def tao_header_trang(doc):

    '''Thêm header với nội dung: Trang {PAGE}/{NUMPAGES} - Mã đề <made>'''

    section = doc.sections[0]

    header = section.header

    paragraph = header.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run('Trang ')

    run.font.color.rgb = RGBColor(0, 0, 255)

    run.font.name = 'Times New Roman'

    run.font.size = Pt(12)

    insert_field_code(paragraph, 'PAGE')

    run = paragraph.add_run('/')

    run.font.name = 'Times New Roman'

    run.font.size = Pt(12)

    insert_field_code(paragraph, 'NUMPAGES')

    run = paragraph.add_run(' - Mã đề <made>')

    run.font.color.rgb = RGBColor(0, 0, 255)

    run.font.name = 'Times New Roman'

    run.font.size = Pt(12)

    set_paragraph_top_border(paragraph, color = 'auto', size = '4', space = '1')





def tao_footer_trang(doc):

    '''Thêm footer với nội dung: Trang {PAGE}/{NUMPAGES} - Mã đề <made>'''

    section = doc.sections[0]

    footer = section.footer

    paragraph = footer.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run('Trang ')

    run.font.color.rgb = RGBColor(0, 0, 0)

    run.font.name = 'Times New Roman'

    run.font.size = Pt(11)

    run.italic = True

    insert_field_code(paragraph, 'PAGE')

    run = paragraph.add_run('/')

    run.font.color.rgb = RGBColor(0, 0, 0)

    run.font.name = 'Times New Roman'

    run.font.size = Pt(11)

    run.italic = True

    insert_field_code(paragraph, 'NUMPAGES')

    run = paragraph.add_run(' - Mã đề <made>')

    run.font.color.rgb = RGBColor(0, 0, 0)

    run.font.name = 'Times New Roman'

    run.font.size = Pt(11)

    run.italic = True

    set_paragraph_top_border(paragraph, color = 'auto', size = '4', space = '1')





def tao_file_mau_A4_Mix():

    doc = Document()

    chinh_le_trang(doc)

    tao_2_bang_header(doc)

    tao_phan_than(doc)

    tao_footer_trang(doc)

    convert_font_to_times_new_roman(doc)

    change_font_size(doc)

    canh_before_after_all(doc)

    canh_before_phan(doc)

    canh_deu_2ben_python_docx(doc)

    name = 'file_mau_A4.docx'

    output_doc = os.path.join(os.getcwd(), name)

    doc.save(output_doc)

    messagebox.showinfo('Thông báo', f'''Đã xong, file tạo ra có tên {name}\n và được lưu cùng thư mục chứa chương trình''')





def tao_file_mau_2c_Mix():

    doc = Document()

    chinh_le_trang(doc)

    tao_header_2c(doc)

    tao_phan_than(doc)

    tao_footer_trang(doc)

    convert_font_to_times_new_roman(doc)

    change_font_size(doc)

    canh_before_after_all(doc)

    canh_before_phan(doc)

    canh_deu_2ben_python_docx(doc)

    name = 'file_mau_2C.docx'

    output_doc = os.path.join(os.getcwd(), name)

    doc.save(output_doc)

    messagebox.showinfo('Thông báo', f'''Đã xong, file tạo ra có tên {name}\n và được lưu cùng thư mục chứa chương trình''')



if __name__ == '__main__':

    tao_file_mau_A4_Mix()

    return None

